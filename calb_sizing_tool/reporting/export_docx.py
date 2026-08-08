# -----------------------------------------------------------------------------
# Personal Open-Source Notice
#
# Copyright (c) 2026 Alex.Zhao. All rights reserved.
#
# This repository is released under the MIT License (see LICENSE file).
# Intended use: learning, evaluation, and engineering reference for Utility-scale
# BESS/ESS sizing and Reporting workflows.
#
# DISCLAIMER: This software is provided "AS IS", without warranty of any kind,
# express or implied. In no event shall the author(s) be liable for any claim,
# damages, or other liability arising from, out of, or in connection with the
# software or the use or other dealings in the software.
#
# NOTE: This is a personal project. It is not an official product or statement
# of any company or organization.
# -----------------------------------------------------------------------------

import datetime
import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# Monkeypatch compatibility: allow xpath(..., namespaces=...) calls on
# docx OXML elements where the underlying xpath implementation does not
# accept the `namespaces` keyword (some docx/lxml versions differ).
try:
    # use the concrete BaseOxmlElement used by CT_* classes
    from docx.oxml.shared import BaseOxmlElement
    from lxml import etree

    _orig_xpath = BaseOxmlElement.xpath

    def _xpath_compat(self, path, *args, **kwargs):
        try:
            return _orig_xpath(self, path, *args, **kwargs)
        except TypeError:
            namespaces = kwargs.get("namespaces")
            if namespaces is not None:
                tree = etree.ElementTree(self)
                return tree.xpath(path, namespaces=namespaces)
            return _orig_xpath(self, path, *args)

    BaseOxmlElement.xpath = _xpath_compat
except Exception:
    pass

from calb_sizing_tool.config import DC_DATA_PATH, PROJECT_ROOT

# ----------------------------------------
# Shared DOCX helpers (match DC report style)
# ----------------------------------------


def _resolve_logo_path() -> Path | None:
    candidates = [
        PROJECT_ROOT / "calb_assets" / "logo" / "calb_logo.png",
        PROJECT_ROOT / "calb_logo.png",
    ]
    for candidate in candidates:
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue
    return None


def _setup_margins(doc: Document):
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)


def add_header_logo(document: Document, logo_path_or_bytes, width=Inches(1.2)) -> list:
    tables = []
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        for para in list(header.paragraphs):
            para._element.getparent().remove(para._element)
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.9))
        tables.append(header_table)
        if logo_path_or_bytes:
            p_logo = header_table.rows[0].cells[0].paragraphs[0]
            run_logo = p_logo.add_run()
            if isinstance(logo_path_or_bytes, (str, Path)):
                run_logo.add_picture(str(logo_path_or_bytes), width=width)
            else:
                run_logo.add_picture(io.BytesIO(logo_path_or_bytes), width=width)
    return tables


def apply_header_logo(document: Document, logo_path=None, width=Inches(1.2)) -> list:
    if logo_path is None:
        logo_path = _resolve_logo_path()
    return add_header_logo(document, logo_path, width=width)


def _apply_footer(doc: Document, footer_lines: list[str] | None) -> None:
    if not footer_lines:
        return
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in list(footer.paragraphs):
            para._element.getparent().remove(para._element)
        p_footer = footer.add_paragraph("\n".join(str(line) for line in footer_lines))
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = p_footer.runs[0] if p_footer.runs else p_footer.add_run()
        run_footer.font.size = Pt(8)


def _add_page_number_footer(doc: Document) -> None:
    """Append a right-aligned 'Page X of Y' field to every section footer.

    Existing footer content is preserved; this adds a new paragraph after it.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _field_run(paragraph, instr: str):
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instr
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_end)
        run.font.size = Pt(8)
        return run

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lead = para.add_run("Page ")
        lead.font.size = Pt(8)
        _field_run(para, "PAGE")
        mid = para.add_run(" of ")
        mid.font.size = Pt(8)
        _field_run(para, "NUMPAGES")


def _setup_header(
    doc: Document,
    title: str = "Confidential Sizing Report",
    *,
    logo_path=None,
    header_lines: list[str] | None = None,
    footer_lines: list[str] | None = None,
) -> None:
    header_tables = apply_header_logo(doc, logo_path=logo_path)

    if header_lines is None:
        header_lines = [
            "CALB Group Co., Ltd.",
            "Utility-Scale Energy Storage Systems",
            f"{title}",
        ]
    else:
        if title and title not in header_lines:
            header_lines = list(header_lines) + [title]

    for header_table in header_tables:
        hdr_cells = header_table.rows[0].cells
        p_info = hdr_cells[1].paragraphs[0]
        p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_info = p_info.add_run("\n".join(str(line) for line in header_lines))
        run_info.font.size = Pt(9)
        run_info.font.bold = True

    _apply_footer(doc, footer_lines)


def _cant_split_row(row) -> None:
    """Prevent a table row from splitting across a page boundary."""
    from docx.oxml import OxmlElement
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _repeat_header_row(row) -> None:
    """Mark a row as a repeating table header (shown at the top of each page)."""
    from docx.oxml import OxmlElement
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def _keep_next_para(para):
    """Keep this paragraph on the same page as the following element. Returns para."""
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))
    return para


def _keep_lines_para(para):
    """Keep all lines of this paragraph together on one page. Returns para."""
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepLines"))
    return para


HEADER_FILL = "1F4E79"  # CALB dark blue
ZEBRA_FILL = "F2F6FA"   # light blue-grey for alternate rows on long tables
_CONTENT_WIDTH_INCHES = 6.9  # US Letter minus 0.8" margins


def _shade_cell(cell, fill_hex: str) -> None:
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill_hex}"/>')
    )


def _add_table(
    doc: Document,
    rows,
    headers,
    *,
    keep_together: bool | None = None,
    col_width_fracs: list[float] | None = None,
    align_right_cols: list[int] | None = None,
):
    if not rows:
        doc.add_paragraph("No data available.")
        return None

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"

    hdr_row = tbl.rows[0]
    hdr = hdr_row.cells
    for j, col in enumerate(headers):
        hdr[j].text = str(col)
    for cell in hdr:
        _shade_cell(cell, HEADER_FILL)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    align_right = set(align_right_cols or [])
    for i, row in enumerate(rows):
        rc = tbl.add_row().cells
        for j, val in enumerate(row):
            rc[j].text = "" if val is None else str(val)
            if j in align_right:
                for para in rc[j].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if len(rows) > 8 and i % 2 == 1:
                _shade_cell(rc[j], ZEBRA_FILL)

    # Label/value two-column tables default to a 62/38 split so values do not
    # float in a half-empty column.
    if col_width_fracs is None and len(headers) == 2:
        col_width_fracs = [0.62, 0.38]
    if col_width_fracs and len(col_width_fracs) == len(headers):
        tbl.autofit = False
        widths = [Inches(_CONTENT_WIDTH_INCHES * frac) for frac in col_width_fracs]
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                cell.width = widths[j]

    # Prevent any row from splitting mid-row across a page boundary
    for row in tbl.rows:
        _cant_split_row(row)

    # Repeat the header row at the top of each page when table is long
    _repeat_header_row(hdr_row)

    # Small tables (≤ 20 data rows): keep the whole table on one page.
    # Large tables: at minimum keep the header row with the first data row.
    should_keep = keep_together if keep_together is not None else (len(rows) <= 20)
    if should_keep:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _keep_next_para(para)
    elif len(tbl.rows) >= 2:
        for cell in hdr_row.cells:
            for para in cell.paragraphs:
                _keep_next_para(para)

    return tbl


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ----------------------------------------
# Cover + Appendix helpers
# ----------------------------------------


def _get_commit_hash():
    try:
        head_path = PROJECT_ROOT / ".git" / "HEAD"
        if not head_path.exists():
            return "unknown"
        head = head_path.read_text(encoding="utf-8").strip()
        if head.startswith("ref:"):
            ref_path = head.split(" ", 1)[1].strip()
            ref_file = PROJECT_ROOT / ".git" / ref_path
            if ref_file.exists():
                return ref_file.read_text(encoding="utf-8").strip()[:7]
        return head[:7]
    except Exception:
        return "unknown"


# ----------------------------------------
# DC dictionary extraction
# ----------------------------------------


# ----------------------------------------
# DC report helper (reuse DC content)
# ----------------------------------------


def _resolve_dc_report_data(dc_output: dict):
    try:
        from calb_sizing_tool.ui import dc_view
    except Exception:
        return None, None, None

    stage1 = dc_output.get("stage1", {}) if dc_output else {}
    if not stage1:
        return None, None, None

    if "results_dict" in dc_output and "report_order" in dc_output:
        return stage1, dc_output.get("results_dict"), dc_output.get("report_order")

    selected = dc_output.get("selected_scenario", stage1.get("selected_scenario", "container_only"))
    try:
        _, df_blocks, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve = dc_view.load_data(DC_DATA_PATH)
        s2, s3_df, s3_meta, iter_count, poi_g, converged = dc_view.size_with_guarantee(
            stage1,
            selected,
            df_blocks,
            df_soh_profile,
            df_soh_curve,
            df_rte_profile,
            df_rte_curve,
            k_max=dc_view.K_MAX_FIXED,
        )
    except Exception:
        return stage1, None, None

    results_dict = {
        selected: (s2, s3_df, s3_meta, iter_count, poi_g, converged)
    }
    report_order = [(selected, selected.replace("_", " ").title())]
    return stage1, results_dict, report_order


def _dc_section_heading(chapter_prefix: str, index: int, title: str):
    if str(chapter_prefix) == "1":
        return f"{index}. {title}"
    return f"{chapter_prefix}.{index} {title}"


def _append_dc_report_sections(doc: Document, dc_output: dict, ctx: dict, chapter_prefix: str = "3"):
    stage1, results_dict, report_order = _resolve_dc_report_data(dc_output or {})
    if stage1 is None or results_dict is None or report_order is None:
        doc.add_paragraph("DC report section unavailable.")
        return False

    try:
        from calb_sizing_tool.ui import dc_view
    except Exception:
        doc.add_paragraph("DC report section unavailable.")
        return False

    heading_level = 2 if str(chapter_prefix) == "1" else 3

    doc.add_heading(
        _dc_section_heading(chapter_prefix, 1, "Project Summary"),
        level=heading_level,
    )
    p = doc.add_paragraph()
    p.add_run(f"Project life: {int(stage1['project_life_years'])} years\n")
    p.add_run(f"POI guarantee year: {int(stage1.get('poi_guarantee_year', 0))}\n")
    p.add_run(f"Cycles per year (assumed): {int(stage1['cycles_per_year'])}\n")
    p.add_run(
        f"S&C time from FAT to COD: {int(round(stage1.get('sc_time_months', 0)))} months\n"
    )
    p.add_run(f"RTE Curve Adjustment (pp): {float(stage1.get('rte_curve_adjust_pp', 0.0)):.1f}\n")
    p.add_run(
        f"DC->POI efficiency chain (one-way): {stage1.get('eff_dc_to_poi_frac', 0.0)*100:.2f}%\n"
    )
    p.add_run(f"POI->DC equivalent power: {stage1.get('dc_power_required_mw', 0.0):.2f} MW")

    doc.add_paragraph(
        "This sizing report is based on the 314 Ah cell database and the internal "
        "CALB SOH/RTE profiles for the selected operating conditions."
    )

    doc.add_heading(
        _dc_section_heading(chapter_prefix, 2, "Equipment Summary (DC Blocks)"),
        level=heading_level,
    )

    for key, title in report_order:
        if key not in results_dict:
            continue
        s2, _, _, iter_count, poi_g, converged = results_dict[key]
        doc.add_paragraph(title, style=None)
        dc_view._docx_add_config_table(doc, s2.get("block_config_table"))
        doc.add_paragraph(f"Iterations: {iter_count} | Guarantee met: {bool(converged)}")
        if poi_g is not None:
            doc.add_paragraph(f"POI usable energy @ guarantee year: {poi_g:.2f} MWh")
        doc.add_paragraph("")

    doc.add_heading(
        _dc_section_heading(
            chapter_prefix,
            3,
            "Lifetime POI Usable Energy & SOH (Per Configuration)",
        ),
        level=heading_level,
    )

    poi_target = float(stage1["poi_energy_req_mwh"])
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))

    for key, title in report_order:
        if key not in results_dict:
            continue
        _, s3_df, s3_meta, _, _, _ = results_dict[key]

        doc.add_paragraph(title, style=None)
        doc.add_paragraph(
            f"POI Power = {s3_meta.get('poi_power_mw', 0.0):.2f} MW | "
            f"DC-equivalent Power = {s3_meta.get('dc_power_mw', 0.0):.2f} MW | "
            f"Effective C-rate (DC-side) = {s3_meta.get('effective_c_rate', 0.0):.3f} C"
        )
        doc.add_paragraph(
            f"SOH profile ID = {s3_meta.get('soh_profile_id')} "
            f"(C-rate <= {s3_meta.get('chosen_soh_c_rate')}, cycles/year = {s3_meta.get('chosen_soh_cycles_per_year')}); "
            f"RTE profile ID = {s3_meta.get('rte_profile_id')} (C-rate <= {s3_meta.get('chosen_rte_c_rate')})."
        )
        doc.add_paragraph(
            f"Guarantee Year (from COD) = {guarantee_year} | POI Energy Target = {poi_target:.2f} MWh"
        )

        if dc_view.MATPLOTLIB_AVAILABLE:
            try:
                cap_png = dc_view._plot_dc_capacity_bar_png(
                    s2=s2,
                    s3_df=s3_df,
                    guarantee_year=guarantee_year,
                    title="DC Block Energy (BOL/COD/Yx at POI)",
                )
                if cap_png and cap_png.getbuffer().nbytes > 0:
                    doc.add_picture(cap_png, width=Inches(6.7))
                png = dc_view._plot_poi_usable_png(
                    s3_df=s3_df,
                    poi_target=poi_target,
                    title=f"POI Usable Energy vs Year \u2013 {key}",
                )
                if png and png.getbuffer().nbytes > 0:
                    doc.add_picture(png, width=Inches(6.7))
            except Exception:
                doc.add_paragraph("Chart export skipped due to plotting error.")
        else:
            doc.add_paragraph("Chart export skipped (matplotlib not available).")

        dc_view._docx_add_lifetime_table(doc, s3_df)
        doc.add_paragraph("")

    return True


# ----------------------------------------
# AC report helpers
# ----------------------------------------


# ----------------------------------------
# Public report generators
# ----------------------------------------


def make_report_filename(proj_name, suffix="Report"):
    safe = "".join(c if c.isalnum() else "_" for c in proj_name)
    return f"{safe}_{suffix}.docx"


def sanitize_filename(text: str, max_length: int = 80) -> str:
    cleaned = re.sub(r'[\\/:*?\"<>|\x00-\x1f]', '', text or '')
    cleaned = re.sub(r'\s+', '_', cleaned.strip())
    cleaned = cleaned.strip('._')
    if max_length and len(cleaned) > max_length:
        cleaned = cleaned[:max_length].rstrip('_')
    return cleaned


def make_proposal_filename(project_name: str | None, version: str = "V2.2",
                           prefix: str = "CALB",
                           ac_alternative: str | None = None) -> str:
    """Proposal file name, optionally naming the AC alternative it reports.

    ``ac_alternative`` is the short label from
    ac_run_service.ac_alternative_label ("A", "B", …). It is appended so two
    alternatives of one DC run download as two files instead of one overwriting
    the other. It is None when the run has only one alternative, so the ordinary
    case keeps the file name it always had.
    """
    stamp = datetime.date.today().strftime("%Y%m%d")
    safe_project = sanitize_filename(project_name or "")
    safe_version = sanitize_filename(version or "", max_length=12) or "V2.2"
    safe_prefix = sanitize_filename(prefix or "", max_length=16) or "CALB"
    safe_alt = sanitize_filename(ac_alternative or "", max_length=8)
    suffix = f"_{safe_version}_AC-{safe_alt}" if safe_alt else f"_{safe_version}"
    if safe_project:
        return f"{safe_prefix}_{safe_project}_BESS_Proposal_{stamp}{suffix}.docx"
    return f"{safe_prefix}_BESS_Proposal_{stamp}{suffix}.docx"


def create_dc_report(dc_output: dict, ctx: dict) -> bytes:
    """REGRESSION BASELINE for the DC page's export — not a product path.

    Nothing in the application calls this. It exists so
    `tests/test_report.py::test_dc_report_unchanged_paragraphs` can hold
    `dc_view.build_report_bytes` — the DOCX behind the DC Sizing page's
    "Export Technical Sizing Report" button, which IS customer-facing — to a
    fixed set of paragraphs. Change this only in step with that page.

    The customer-facing proposal is `report_v2.export_report_v2_1`, reached
    from the Report Export page. Owner ruling 2026-08-08 kept this baseline and
    deleted the AC and combined generators, which anchored nothing.
    """
    doc = Document()
    _setup_margins(doc)
    _setup_header(doc)

    stage1 = dc_output.get("stage1", {}) if dc_output else {}
    project_name = stage1.get("project_name", "CALB ESS Project")

    doc.add_heading("CALB Utility-Scale ESS Sizing Report", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(
        f"POI Requirement: {stage1['poi_power_req_mw']:.2f} MW / "
        f"{stage1['poi_energy_req_mwh']:.2f} MWh"
    )
    doc.add_paragraph("")

    _append_dc_report_sections(doc, dc_output, ctx, chapter_prefix="1")
    return _doc_to_bytes(doc)


