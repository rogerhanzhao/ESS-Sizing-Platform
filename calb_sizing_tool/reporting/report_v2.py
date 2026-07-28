# -----------------------------------------------------------------------------
# Personal Open-Source Notice
#
# Copyright (c) 2026 Alex.Zhao. All rights reserved.
#
# This repository is released under the MIT License (see LICENSE file).
# Intended use: learning, evaluation, and engineering reference for Utility-scale
# BESS/ESS sizing and Reporting workflows.
#
# DISCLAIMER: This software is provided "AS IS", without warranty of any kind,
# express or implied. In no event shall the author(s) be liable for any claim,
# damages, or other liability arising from, out of, or in connection with the
# software or the use or other dealings in the software.
#
# NOTE: This is a personal project. It is not an official product or statement
# of any company or organization.
# -----------------------------------------------------------------------------

import datetime
import io
import re
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Inches

from calb_sizing_tool.reporting.export_docx import (
    _add_page_number_footer,
    _add_table,
    _doc_to_bytes,
    _keep_next_para,
    _setup_header,
    _setup_margins,
)
from calb_diagrams.ac_block_arrangement_v2 import (
    US_NFPA_OIL as ARRANGEMENT_PROFILE,
    render_plan_svg as render_ac_block_plan_svg,
)
from calb_diagrams.ac_block_bilateral_layout import (
    LAYOUT_VARIANT as BILATERAL_LAYOUT_VARIANT,
    compute_bilateral_layout,
    render_bilateral_plan_svg,
)
from calb_diagrams.governed_site_layout_concept import render_governed_site_layout_concept_svg
from calb_diagrams.site_array_concept import (
    US_NFPA_SITE as SITE_PROFILE,
    compute_site_array,
    render_site_svg,
)
from calb_sizing_tool.reporting.brand_profiles import (
    BrandProfile,
    CALB_BRAND,
    neutralize_equipment_text,
    require_brand_assets,
)
from calb_sizing_tool.reporting.formatter import format_percent, format_value
from calb_sizing_tool.reporting.report_context import ReportContext

try:
    from matplotlib.figure import Figure

    MATPLOTLIB_AVAILABLE = True
except Exception:
    MATPLOTLIB_AVAILABLE = False

try:
    import cairosvg

    CAIROSVG_AVAILABLE = True
except Exception:
    CAIROSVG_AVAILABLE = False

from calb_sizing_tool.common.render_lock import RENDER_LOCK


def _format_percent_with_fraction(value, input_is_fraction=None, fraction_decimals=4) -> str:
    """Legacy helper retained for compatibility; prefer format_percent() in new code."""
    if value is None:
        return ""
    try:
        numeric = float(value)
    except Exception:
        return str(value)

    if input_is_fraction is None:
        is_fraction = numeric <= 1.2
    else:
        is_fraction = bool(input_is_fraction)

    fraction = numeric if is_fraction else numeric / 100.0
    percent_text = format_percent(numeric, input_is_fraction=is_fraction)
    return f"{percent_text} ({fraction:.{fraction_decimals}f})"


def _default_formatter(value):
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return str(value)


def _add_dataframe_table(doc: Document, df: Optional[pd.DataFrame], columns, headers_map, formatters, *, keep_together: bool | None = None):
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return

    rows = []
    for _, row in df.iterrows():
        rows.append([formatters.get(col, _default_formatter)(row.get(col)) for col in columns])

    headers = [headers_map.get(col, col) for col in columns]
    numeric_cols = [
        j for j, col in enumerate(columns)
        if col in df.columns and pd.api.types.is_numeric_dtype(df[col])
    ]
    _add_table(doc, rows, headers, keep_together=keep_together, align_right_cols=numeric_cols)


def _plot_poi_usable_png(df: pd.DataFrame, poi_target: float, title: str) -> Optional[io.BytesIO]:
    if not MATPLOTLIB_AVAILABLE:
        return None
    try:
        data = df.sort_values("Year_Index").copy()
        x = data["Year_Index"].astype(int).tolist()
        y = data["POI_Usable_Energy_MWh"].astype(float).tolist()

        fig = Figure(figsize=(7.0, 3.2))
        ax = fig.add_subplot(111)
        ax.bar(x, y, color="#5cc3e4")
        ax.axhline(poi_target, linewidth=2, color="#ff0000")
        ax.set_title(title)
        ax.set_xlabel("Year (from COD)")
        ax.set_ylabel("POI Usable Energy (MWh)")
        ax.set_xticks(x)
        # Headroom above the tallest bar / target line so the red guarantee
        # line never sits on the plot border.
        top = max(max(y) if y else 0.0, float(poi_target or 0.0))
        if top > 0:
            ax.set_ylim(0, top * 1.12)
        ax.grid(True, axis="y", linestyle="--", alpha=0.35)

        buf = io.BytesIO()
        fig.tight_layout()
        with RENDER_LOCK:
            fig.savefig(buf, format="png", dpi=150)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _plot_dc_capacity_bar_png(
    bol_mwh: Optional[float],
    s3_df: Optional[pd.DataFrame],
    guarantee_year: int,
    poi_target: Optional[float],
    title: str,
) -> Optional[io.BytesIO]:
    """Bar chart: installed DC nameplate vs deliverable POI energy at key milestones.

    BOL is a DC-side metric; all other bars are POI-side. The two measurement
    planes use distinct colors and an explicit legend so the DC/POI distinction
    is unambiguous, and the guarantee target line is drawn across the POI bars
    only (the target does not apply to the DC plane).
    """
    if not MATPLOTLIB_AVAILABLE:
        return None
    try:
        from matplotlib.lines import Line2D
        from matplotlib.patches import Patch

        if s3_df is None or s3_df.empty or bol_mwh is None:
            return None

        def _poi_at(year: int) -> Optional[float]:
            row = s3_df[s3_df["Year_Index"] == int(year)]
            if row.empty:
                return None
            return float(row["POI_Usable_Energy_MWh"].iloc[0])

        g_year = int(guarantee_year)
        final_year = int(s3_df["Year_Index"].max())

        DC_COLOR = "#1f4e79"
        POI_COLOR = "#5cc3e4"

        labels = ["DC Nameplate\nBOL (DC side)"]
        values = [float(bol_mwh)]
        colors = [DC_COLOR]

        cod = _poi_at(0)
        if cod is not None:
            cod_label = "POI Usable\nY0 (COD)"
            if g_year == 0:
                cod_label = "POI Usable\nY0 (COD, guarantee)"
            labels.append(cod_label)
            values.append(cod)
            colors.append(POI_COLOR)

        if g_year > 0:
            yg = _poi_at(g_year)
            if yg is not None:
                labels.append(f"POI Usable\nY{g_year} (guarantee)")
                values.append(yg)
                colors.append(POI_COLOR)

        if final_year not in (0, g_year):
            y_end = _poi_at(final_year)
            if y_end is not None:
                labels.append(f"POI Usable\nY{final_year} (end of life)")
                values.append(y_end)
                colors.append(POI_COLOR)

        fig = Figure(figsize=(7.0, 3.4))
        ax = fig.add_subplot(111)
        bars = ax.bar(labels, values, color=colors, width=0.6)
        ax.bar_label(bars, fmt="%.1f", padding=2, fontsize=9)

        legend_handles = [
            Patch(facecolor=DC_COLOR, label="DC side (installed nameplate)"),
            Patch(facecolor=POI_COLOR, label="POI side (deliverable energy)"),
        ]
        if poi_target is not None and float(poi_target) > 0 and len(values) > 1:
            ax.hlines(
                float(poi_target),
                xmin=0.7,
                xmax=len(values) - 1 + 0.3,
                linewidth=2,
                color="#ff0000",
            )
            legend_handles.append(
                Line2D([0], [0], color="#ff0000", linewidth=2, label="POI guarantee target")
            )

        ax.set_title(title)
        ax.set_ylabel("Energy (MWh)")
        ax.set_ylim(0, max(values) * 1.18)
        ax.grid(True, axis="y", linestyle="--", alpha=0.35)
        ax.legend(handles=legend_handles, loc="upper right", fontsize=8, framealpha=0.9)

        buf = io.BytesIO()
        fig.tight_layout()
        with RENDER_LOCK:
            fig.savefig(buf, format="png", dpi=150)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _add_cover_page_v2(doc: Document, ctx: ReportContext, brand: BrandProfile) -> None:
    """Proposal cover: title, project identity block, issuer, confidentiality note.

    All branded copy comes from the BrandProfile — no fallback strings here.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    for _ in range(4):
        doc.add_paragraph("")

    heading = doc.add_heading(brand.cover_title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Battery Energy Storage System — Sizing Proposal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle.runs:
        subtitle.runs[0].font.size = Pt(13)

    doc.add_paragraph("")
    generated = ctx.report_generated_at or datetime.datetime.now().strftime("%Y-%m-%d")
    info_lines = [
        f"Project: {ctx.project_name}",
        f"Case: {ctx.case_name or '—'}",
        f"Date: {generated}",
        f"Tool Version: {brand.tool_version_label}",
    ]
    info = doc.add_paragraph("\n".join(info_lines))
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    issuer = doc.add_paragraph("\n".join(brand.issuer_lines))
    issuer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(8):
        doc.add_paragraph("")
    notice = doc.add_paragraph(brand.confidentiality_notice)
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in notice.runs:
        run.font.size = Pt(8)
        run.italic = True

    doc.add_page_break()


def _mixed_head_entry(ctx: ReportContext):
    """The Head AC Block model entry of a genuinely mixed station, else None.

    Uniform stations (a single AC Block model, however DC Blocks are
    distributed) carry a single-entry breakdown and return None, so callers fall
    back to the plain average.
    """
    ac_output = ctx.ac_output if isinstance(ctx.ac_output, dict) else {}
    if not ac_output.get("ac_block_mixed"):
        return None
    breakdown = ac_output.get("ac_block_breakdown")
    if not isinstance(breakdown, list) or len(breakdown) <= 1:
        return None
    return next((entry for entry in breakdown if entry.get("role") == "Head"), breakdown[0])


def _representative_dc_per_ac(ctx: ReportContext) -> int:
    """DC Blocks per AC Block for the *representative* block drawn in §8/§9.

    A mixed station has no single average block — the fractional mean matches no
    real AC Block. Draw the HEAD block instead (consistent with the §6.1
    schedule and the SLD head-fleet projection); tail blocks are described in
    §6.1. Uniform stations keep the plain average.
    """
    head = _mixed_head_entry(ctx)
    if head is not None:
        dc_each = head.get("dc_blocks_each")
        if dc_each is None:
            dc_each = head.get("dc_blocks_max") or head.get("dc_blocks_min")
        return max(1, int(dc_each or 1))
    if ctx.ac_blocks_total <= 0:
        return 0
    return max(1, int(round(ctx.dc_blocks_total / ctx.ac_blocks_total)))


def _compute_site_layout(ctx: ReportContext):
    """Concept site-array layout from the sizing result (None if not sizeable)."""
    if ctx.ac_blocks_total <= 0 or ctx.dc_blocks_total <= 0:
        return None
    # The linear L2 site-array engine would draw single-row DC fields, which
    # contradicts a governed bilateral 4+4 unit and its SLD. Whole-site
    # composition of bilateral units is Master Layout (L3) scope, so suppress
    # the linear site figure here rather than render inconsistent geometry.
    if ctx.layout_variant == BILATERAL_LAYOUT_VARIANT:
        return None
    dc_per_ac = _representative_dc_per_ac(ctx)
    if dc_per_ac <= 0:
        return None
    try:
        return compute_site_array(
            ctx.ac_blocks_total, dc_per_ac, ARRANGEMENT_PROFILE, SITE_PROFILE
        )
    except Exception:
        return None


def _compute_typical_group_layout(ctx: ReportContext):
    """A single representative project group, drawn at legible page scale.

    The full-site concept array is a tall linear strip (all groups stacked in
    one column), which collapses to an illegible sliver when fit to a page. A
    real site simply repeats one project group, so the report draws ONE typical
    group at legible scale and states the whole-site composition (N such groups)
    in the summary instead of drawing every block. For a mixed station the
    representative block is the Head AC Block (tails are listed in §6.1).
    """
    if ctx.ac_blocks_total <= 0 or ctx.dc_blocks_total <= 0:
        return None
    if ctx.layout_variant == BILATERAL_LAYOUT_VARIANT:
        return None
    dc_per_ac = _representative_dc_per_ac(ctx)
    if dc_per_ac <= 0:
        return None
    group_blocks = min(SITE_PROFILE.default_blocks_per_group, ctx.ac_blocks_total)
    try:
        return compute_site_array(
            group_blocks, dc_per_ac, ARRANGEMENT_PROFILE, SITE_PROFILE
        )
    except Exception:
        return None


def _governed_run_from_ctx(ctx: ReportContext):
    """Assemble a governed site run-like object from the report context.

    Geometry / grouping comes from the deterministic decomposition of the DC
    total; the per-group transformer nameplate, vector group and bound product
    come from the ACTUAL AC Sizing output on ``ctx`` (never re-bound here), so the
    site layout and equipment schedule match what the AC Sizing page produced. A
    mixed run carries this per group in ``governed_groups``; a pure run carries it
    once on the output itself.
    """
    from types import SimpleNamespace

    from calb_sizing_tool.schemas.governed_ac_block_config import (
        get_governed_configuration,
    )
    from calb_sizing_tool.services.governed_ac_block_service import (
        build_governed_site_plan,
    )

    ac_output = ctx.ac_output if isinstance(ctx.ac_output, dict) else {}

    # Preferred: consume the persisted, versioned site run recorded at AC Sizing
    # time — the report never re-decomposes, so it cannot drift if the catalogue
    # or decomposition rules change after the run was saved.
    persisted = ac_output.get("governed_site_run")
    if isinstance(persisted, dict) and persisted.get("groups"):
        groups = []
        for gg in persisted["groups"]:
            groups.append(SimpleNamespace(
                configuration_code=gg.get("configuration_code"),
                layout_variant=gg.get("layout_variant"),
                ac_block_count=int(gg.get("ac_block_count") or 0),
                dc_blocks_per_ac_block=int(gg.get("dc_blocks_per_ac_block") or 0),
                pcs_per_ac_block=int(gg.get("pcs_per_ac_block") or 0),
                pcs_kw=float(gg.get("pcs_kw") or 0.0),
                ac_power_mw_total=float(gg.get("ac_power_mw_total") or 0.0),
                bound_product_code=gg.get("bound_product_code"),
                product_confirmation=gg.get("product_confirmation") or "none",
                ac_output={
                    "transformer_mva": gg.get("transformer_mva"),
                    "transformer_vector_group": gg.get("transformer_vector_group"),
                },
            ))
        return SimpleNamespace(
            dc_blocks_total=int(persisted.get("dc_blocks_total") or 0),
            ac_blocks_total=int(persisted.get("ac_blocks_total") or 0),
            ac_power_mw_total=float(persisted.get("ac_power_mw_total") or 0.0),
            groups=groups,
        )

    # Legacy fallback: an older run without a persisted site run — reconstruct the
    # geometry deterministically from the DC total (as before).
    plan = build_governed_site_plan(int(ctx.dc_blocks_total or 0))
    governed_groups = ac_output.get("governed_groups")

    # code -> (transformer_mva, vector_group, product, product_confirmation)
    by_code: dict[str, tuple] = {}
    if isinstance(governed_groups, list) and governed_groups:
        for gg in governed_groups:
            by_code[str(gg.get("configuration_code"))] = (
                gg.get("transformer_mva"),
                None,  # tail vector groups are not carried per group
                gg.get("bound_product_code"),
                gg.get("product_confirmation") or ("owner_selected" if gg.get("bound_product_code") else "none"),
            )
        # the head also carries a confirmed vector group on the output
        head_code = str(ac_output.get("configuration_code") or "")
        if head_code in by_code:
            mva, _vg, prod, conf = by_code[head_code]
            by_code[head_code] = (mva, ac_output.get("transformer_vector_group"), prod, conf)
    elif ac_output.get("configuration_code"):
        prod = ac_output.get("governed_product_block_code")
        by_code[str(ac_output.get("configuration_code"))] = (
            ac_output.get("transformer_mva"),
            ac_output.get("transformer_vector_group"),
            prod,
            "owner_selected" if prod else "none",
        )

    groups = []
    for g in plan.groups:
        mva, vg, prod, conf = by_code.get(g.configuration_code, (None, None, None, "none"))
        try:
            pcs_kw = float(get_governed_configuration(g.configuration_code).pcs_rating_kw)
        except Exception:
            pcs_kw = 0.0
        groups.append(SimpleNamespace(
            configuration_code=g.configuration_code,
            layout_variant=g.layout_variant,
            ac_block_count=g.ac_block_count,
            dc_blocks_per_ac_block=g.dc_blocks_per_ac_block,
            pcs_per_ac_block=g.pcs_per_ac_block,
            pcs_kw=pcs_kw,
            ac_power_mw_total=g.ac_power_mw_total,
            bound_product_code=prod,
            product_confirmation=conf,
            ac_output={"transformer_mva": mva, "transformer_vector_group": vg},
        ))
    return SimpleNamespace(
        dc_blocks_total=plan.dc_blocks_total,
        ac_blocks_total=plan.ac_blocks_total,
        ac_power_mw_total=plan.ac_power_mw_total,
        groups=groups,
    )


def _normalize_template_label(value) -> str:
    """Normalize AC block template ids like '4x1250kw' to '4×1250 kW' for display."""
    if value is None:
        return "—"
    text = str(value)
    match = re.fullmatch(r"\s*(\d+)\s*[xX×]\s*(\d+)\s*[kK][wW]\s*", text)
    if match:
        return f"{match.group(1)}×{match.group(2)} kW"
    return text


def _format_or_tbd(value, unit: str) -> str:
    if value is None:
        return "TBD"
    try:
        numeric = float(value)
    except Exception:
        return str(value) if value else "TBD"
    if numeric <= 0:
        return "TBD"
    return format_value(numeric, unit)


def _format_transformer_rating(value) -> str:
    if value is None:
        return "TBD"
    try:
        kva = float(value)
    except Exception:
        return "TBD"
    if kva <= 0:
        return "TBD"
    mva = kva / 1000.0
    return f"{mva:.2f} MVA ({kva:.0f} kVA)"


def _svg_bytes_to_png(svg_bytes: bytes, width_px: int = 900) -> Optional[bytes]:
    if not svg_bytes or not CAIROSVG_AVAILABLE:
        return None
    try:
        with RENDER_LOCK:
            return cairosvg.svg2png(bytestring=svg_bytes, output_width=width_px)
    except Exception:
        return None


# Governance rule (owner decision, 2026-07-28): every concept engineering figure
# in the exported report — the SLD (§7), the Typical AC Block Arrangement (§8)
# and the Concept Site Layout / Arrangement (§9) — is stamped
# "DRAFT / OVERRIDE - NOT FOR CONSTRUCTION" UNCONDITIONALLY, no matter how
# professional the drawing is or how the layout resolves. The report is a
# concept / proposal document and is never a construction-issue drawing set.
# The text, colour and format match the existing SLD-pipeline watermark
# (_concept_safe_svg): #B42318 at 0.28 opacity, bold, horizontal-centred.
NOT_FOR_CONSTRUCTION_STAMP = "DRAFT / OVERRIDE - NOT FOR CONSTRUCTION"
_STAMP_FILL = (180, 35, 24, 71)  # #B42318 @ ~0.28 opacity
_STAMP_RED = (180, 35, 24)  # #B42318 opaque — placeholder border/text


class WatermarkError(RuntimeError):
    """The mandatory NOT-FOR-CONSTRUCTION mark could not be applied.

    Raised only when the figure can neither be watermarked nor replaced with a
    visibly-marked placeholder (e.g. Pillow itself is unavailable). Enforcement
    is fail-closed: the report aborts rather than embed an unmarked drawing.
    """


def _load_stamp_font(draw, text: str, width: int) -> "object":
    """Pick a bold font sized so *text* spans ~82% of *width* (shared by the
    watermark and its failure placeholder)."""
    from PIL import ImageFont

    def _load(size: int):
        for candidate in ("DejaVuSans-Bold.ttf", "DejaVuSans.ttf"):
            try:
                return ImageFont.truetype(candidate, size)
            except Exception:
                continue
        return ImageFont.load_default()

    ref_size = 100
    ref_font = _load(ref_size)
    ref_bbox = draw.textbbox((0, 0), text, font=ref_font)
    ref_w = max(1, ref_bbox[2] - ref_bbox[0])
    font_size = max(14, int(ref_size * (width * 0.82) / ref_w))
    return _load(font_size)


def _watermark_failure_placeholder(png_bytes: Optional[bytes]) -> Optional[bytes]:
    """Opaque red-bordered placeholder embedded when watermarking fails.

    Deliberately does NOT contain the source drawing — the whole point of
    fail-closed enforcement is that an un-stampable figure must never reach the
    report as a clean engineering drawing. Returns ``None`` only when even a
    placeholder cannot be produced (Pillow unavailable), so the caller can raise.
    """
    try:
        from PIL import Image, ImageDraw

        width, height = 1600, 900
        try:
            with Image.open(io.BytesIO(png_bytes)) as src:
                width, height = src.size
        except Exception:
            pass  # corrupt/unknown source — fall back to a default canvas size

        img = Image.new("RGB", (width, height), (255, 240, 238))  # pale red wash
        draw = ImageDraw.Draw(img)
        border = max(4, min(width, height) // 60)
        draw.rectangle([0, 0, width - 1, height - 1], outline=_STAMP_RED, width=border)

        lines = [NOT_FOR_CONSTRUCTION_STAMP, "FIGURE WITHHELD — watermark could not be applied"]
        font = _load_stamp_font(draw, lines[1], width)
        y = height * 0.5
        for i, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
            draw.text(((width - tw) / 2.0 - bbox[0],
                       y + (i - 0.5) * th * 1.6 - bbox[1]),
                      line, font=font, fill=_STAMP_RED)
        out = io.BytesIO()
        img.save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def _stamp_not_for_construction(png_bytes: Optional[bytes]) -> Optional[bytes]:
    """Overlay the translucent red "DRAFT / OVERRIDE - NOT FOR CONSTRUCTION" mark.

    Matches the SLD-pipeline watermark style (horizontal, centred, #B42318 @ 0.28)
    and is applied to every concept figure before it is embedded, so the mark is
    present regardless of the figure's own document status.

    Fail-closed: if the source image cannot be stamped, a visibly-marked red
    placeholder is returned instead of the original — an un-stampable figure must
    never be embedded as a clean drawing. If even the placeholder cannot be built
    (Pillow unavailable) a :class:`WatermarkError` is raised so the report aborts
    rather than silently emit an unmarked figure.
    """
    if not png_bytes:
        return png_bytes
    try:
        from PIL import Image, ImageDraw

        base = Image.open(io.BytesIO(png_bytes)).convert("RGBA")
        width, height = base.size
        text = NOT_FOR_CONSTRUCTION_STAMP
        layer = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        draw = ImageDraw.Draw(layer)

        # Scale the mark so the text spans ~82% of the figure width (matches the
        # moderate SVG overlay) and always fits, regardless of source resolution.
        font = _load_stamp_font(draw, text, width)

        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        # Horizontal, centred at ~52% height — same placement as the SVG overlay.
        draw.text(((width - tw) / 2.0 - bbox[0], height * 0.52 - th / 2.0 - bbox[1]),
                  text, font=font, fill=_STAMP_FILL)
        stamped = Image.alpha_composite(base, layer).convert("RGB")
        out = io.BytesIO()
        stamped.save(out, format="PNG")
        return out.getvalue()
    except Exception as exc:
        placeholder = _watermark_failure_placeholder(png_bytes)
        if placeholder is not None:
            return placeholder
        raise WatermarkError(
            "Could not apply the mandatory NOT-FOR-CONSTRUCTION watermark and "
            "could not build a placeholder; refusing to embed an unmarked figure."
        ) from exc


def _add_concept_figure(doc, png_bytes: bytes, *, width=None, height=None) -> None:
    """Embed a concept figure with the mandatory NOT-FOR-CONSTRUCTION watermark."""
    stamped = _stamp_not_for_construction(png_bytes)
    if width is not None:
        doc.add_picture(io.BytesIO(stamped), width=width)
    elif height is not None:
        doc.add_picture(io.BytesIO(stamped), height=height)
    else:
        doc.add_picture(io.BytesIO(stamped))


def _validate_efficiency_chain(ctx: ReportContext) -> list[str]:
    """Validate efficiency chain completeness and internal consistency.

    Returns advisory warning messages; does not block export.
    """
    warnings = []

    if not isinstance(ctx.stage1, dict) or not ctx.stage1:
        warnings.append(
            "Cannot validate efficiency: stage1 (DC SIZING output) is missing or invalid. "
            "Ensure DC SIZING was completed before exporting."
        )

    components = [
        ("DC Cables", ctx.efficiency_components_frac.get("eff_dc_cables_frac")),
        ("PCS", ctx.efficiency_components_frac.get("eff_pcs_frac")),
        ("Transformer (MVT)", ctx.efficiency_components_frac.get("eff_mvt_frac")),
        ("RMU/Switchgear/AC Cables", ctx.efficiency_components_frac.get("eff_ac_cables_sw_rmu_frac")),
        ("HVT/Others", ctx.efficiency_components_frac.get("eff_hvt_others_frac")),
    ]

    has_all_components = True
    for name, value in components:
        if value is None or value <= 0:
            warnings.append(f"Efficiency component '{name}' is missing or zero: {value}.")
            has_all_components = False
        elif value > 1.2:
            warnings.append(f"Efficiency component '{name}' exceeds 120%: {value}")

    if ctx.efficiency_chain_oneway_frac is None or ctx.efficiency_chain_oneway_frac <= 0:
        warnings.append(f"Total one-way efficiency chain is missing or zero: {ctx.efficiency_chain_oneway_frac}.")
    elif ctx.efficiency_chain_oneway_frac > 1.2:
        warnings.append(f"Total one-way efficiency exceeds 120%: {ctx.efficiency_chain_oneway_frac}")

    if has_all_components and ctx.efficiency_chain_oneway_frac and ctx.efficiency_chain_oneway_frac > 0:
        product = 1.0
        for name, value in components:
            if value and value > 0:
                product *= value
        relative_error = abs(product - ctx.efficiency_chain_oneway_frac) / ctx.efficiency_chain_oneway_frac
        if relative_error > 0.02:
            warnings.append(
                f"Total efficiency ({ctx.efficiency_chain_oneway_frac:.6f}) "
                f"does not match product of components ({product:.6f}); "
                f"relative error: {relative_error * 100:.2f}%."
            )

    if ctx.efficiency_chain_oneway_frac is not None and ctx.efficiency_chain_oneway_frac < 0.001:
        warnings.append("Efficiency chain appears uninitialized; ensure DC SIZING was completed.")

    return warnings


def _aggregate_ac_block_configs(ctx: ReportContext) -> list[dict]:
    """Aggregate AC Block configurations by signature (PCS count, rating, power per block).

    Returns list of dicts with keys: pcs_per_block, pcs_kw, ac_block_power_mw, count.
    """
    if ctx.ac_blocks_total == 0:
        return []

    pcs_per_block = ctx.pcs_per_block
    pcs_kw = None
    if isinstance(ctx.ac_output, dict) and ctx.ac_output.get("pcs_kw"):
        pcs_kw = ctx.ac_output.get("pcs_kw")
    if pcs_kw is None and isinstance(ctx.ac_output, dict):
        pcs_kw = ctx.ac_output.get("pcs_power_kw")

    ac_block_power_mw = ctx.ac_block_size_mw
    if pcs_kw is None and ac_block_power_mw and pcs_per_block and pcs_per_block > 0:
        pcs_kw = int((ac_block_power_mw * 1000) / pcs_per_block)

    return [
        {
            "pcs_per_block": pcs_per_block,
            "pcs_kw": pcs_kw,
            "ac_block_power_mw": ac_block_power_mw,
            "count": ctx.ac_blocks_total,
        }
    ]


def _validate_report_consistency(ctx: ReportContext) -> list[str]:
    """Validate overall report consistency (power/energy/efficiency).

    Returns advisory warning messages; does not block export.
    """
    warnings = []

    warnings.extend(_validate_efficiency_chain(ctx))

    if ctx.ac_blocks_total > 0 and ctx.dc_blocks_total == 0:
        warnings.append("AC Blocks present but DC Blocks count is zero.")

    expected_pcs = ctx.ac_blocks_total * ctx.pcs_per_block
    if ctx.pcs_modules_total > 0 and ctx.pcs_modules_total != expected_pcs:
        warnings.append(
            f"PCS module count mismatch: expected {expected_pcs} "
            f"(AC blocks={ctx.ac_blocks_total} x PCS/block={ctx.pcs_per_block}), "
            f"got {ctx.pcs_modules_total}."
        )

    if ctx.ac_blocks_total > 0 and ctx.ac_block_size_mw and ctx.ac_block_size_mw > 0:
        total_ac_power = ctx.ac_blocks_total * ctx.ac_block_size_mw
        poi_requirement = ctx.poi_power_requirement_mw
        overage = total_ac_power - poi_requirement
        overage_pct = (overage / poi_requirement * 100) if poi_requirement > 0 else 0
        if overage > 0.5 and overage_pct > 10:
            warnings.append(
                f"AC power overbuild is {overage_pct:.1f}% "
                f"(total {total_ac_power:.2f} MW vs requirement {poi_requirement:.2f} MW)."
            )

    if ctx.dc_total_energy_mwh is not None and ctx.poi_energy_requirement_mwh is not None:
        if ctx.dc_total_energy_mwh < ctx.poi_energy_requirement_mwh:
            warnings.append(
                f"DC nameplate ({ctx.dc_total_energy_mwh:.2f} MWh) is less than "
                f"POI requirement ({ctx.poi_energy_requirement_mwh:.2f} MWh); "
                f"degradation modeling determines actual delivery."
            )

    if (ctx.poi_usable_energy_mwh_at_guarantee_year is not None
            and ctx.poi_energy_guarantee_mwh is not None):
        if ctx.poi_usable_energy_mwh_at_guarantee_year + 0.1 < ctx.poi_energy_guarantee_mwh:
            warnings.append(
                f"POI usable at guarantee year ({ctx.poi_usable_energy_mwh_at_guarantee_year:.2f} MWh) "
                f"is below guarantee target ({ctx.poi_energy_guarantee_mwh:.2f} MWh)."
            )

    if ctx.poi_guarantee_year > ctx.project_life_years:
        warnings.append(
            f"Guarantee year ({ctx.poi_guarantee_year}) exceeds project life ({ctx.project_life_years} years)."
        )

    return warnings


def export_report_v2_1(ctx: ReportContext, brand: BrandProfile | None = None) -> bytes:
    if brand is None:
        brand = CALB_BRAND
    require_brand_assets(brand)

    doc = Document()
    _setup_margins(doc)
    _setup_header(
        doc,
        title=brand.header_title,
        logo_path=brand.logo_path,
        header_lines=list(brand.header_lines),
        footer_lines=list(brand.footer_lines) or None,
    )

    _add_page_number_footer(doc)
    _add_cover_page_v2(doc, ctx, brand)

    # Shared derived values used across multiple sections
    gyr_target = ctx.poi_energy_guarantee_mwh
    poi_usable_at_gyr = ctx.poi_usable_energy_mwh_at_guarantee_year
    meets_guarantee = (
        "Yes"
        if (poi_usable_at_gyr is not None and gyr_target is not None
            and poi_usable_at_gyr >= gyr_target - 0.1)
        else "No"
    )

    # --- Document Provenance (DB linkage) ---
    # _add_cover_page already ends with a page break; no second break needed here.
    doc.add_heading("Document Provenance", level=2)
    dc_source = "Saved (database)" if ctx.run_id else "Live session (unsaved)"
    ac_source = (
        "Saved (database)" if ctx.ac_run_id
        else ("Linked to DC run" if ctx.run_id else "Live session (unsaved)")
    )
    prov_rows = [
        ("Project", ctx.project_name),
        ("Case", ctx.case_name or "—"),
        ("DC Sizing Data", dc_source),
        ("AC Sizing Data", ac_source),
        ("Report Generated", ctx.report_generated_at or datetime.datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    _add_table(doc, prov_rows, ["Field", "Value"])
    if not ctx.run_id:
        doc.add_paragraph(
            "NOTE: Sizing data has not been saved to the database. "
            "Save the DC run before issuing this report for full traceability."
        )

    # --- Section 1: Executive Summary ---
    # Flows on the same page as Document Provenance (no forced page break).
    doc.add_heading("1.  Executive Summary", level=2)
    if poi_usable_at_gyr is not None and gyr_target is not None:
        key_para = doc.add_paragraph()
        key_run = key_para.add_run(
            f"The proposed system delivers {format_value(poi_usable_at_gyr, 'MWh')} MWh at the POI "
            f"in Year {ctx.poi_guarantee_year} against a guarantee target of "
            f"{format_value(gyr_target, 'MWh')} MWh — guarantee "
            f"{'met' if meets_guarantee == 'Yes' else 'NOT met'}."
        )
        key_run.bold = True
        _exec_governed_groups = (
            ctx.ac_output.get("governed_groups") if isinstance(ctx.ac_output, dict) else None
        )
        if isinstance(_exec_governed_groups, list) and _exec_governed_groups:
            # Mixed governed (Phase B): state the true governed decomposition, not
            # a single uniform template or an average DC-per-AC ratio.
            _comp = " + ".join(
                f"{int(g.get('ac_block_count') or 0)} × {g.get('configuration_code')}"
                for g in _exec_governed_groups
            )
            config_para = doc.add_paragraph(
                f"Configuration: {ctx.dc_blocks_total} DC Blocks "
                f"({format_value(ctx.dc_total_energy_mwh, 'MWh')} MWh nameplate @BOL) · "
                f"{ctx.ac_blocks_total} AC Blocks · {ctx.pcs_modules_total} PCS modules — "
                f"{_comp} (see Section 9 for the equipment schedule)."
            )
        else:
            config_para = doc.add_paragraph(
                f"Configuration: {ctx.dc_blocks_total} DC Blocks "
                f"({format_value(ctx.dc_total_energy_mwh, 'MWh')} MWh nameplate @BOL) · "
                f"{ctx.ac_blocks_total} AC Blocks ({_normalize_template_label(ctx.ac_block_template_id)}) · "
                f"{ctx.pcs_modules_total} PCS modules."
            )
        _keep_next_para(key_para)
        _keep_next_para(config_para)
    exec_rows = [
        ("POI Power Requirement (MW)", format_value(ctx.poi_power_requirement_mw, "MW")),
        ("POI Energy Requirement (MWh)", format_value(ctx.poi_energy_requirement_mwh, "MWh")),
        ("POI Energy Guarantee Target (MWh)", format_value(gyr_target, "MWh")),
        ("Guarantee Year (from COD)", f"{ctx.poi_guarantee_year:d}"),
        ("POI Usable @ Guarantee Year (MWh)", format_value(poi_usable_at_gyr, "MWh")),
        ("Guarantee Compliance", meets_guarantee),
        ("DC Blocks Total", f"{ctx.dc_blocks_total:d}"),
        ("DC Nameplate Energy @BOL (MWh)", format_value(ctx.dc_total_energy_mwh, "MWh")),
        ("AC Block Template", _normalize_template_label(ctx.ac_block_template_id)),
        ("AC Blocks Total", f"{ctx.ac_blocks_total:d}"),
        ("Total PCS Modules", f"{ctx.pcs_modules_total:d}"),
        ("Transformer Rating", _format_transformer_rating(ctx.transformer_rating_kva)),
    ]
    site_layout = _compute_site_layout(ctx)
    if site_layout is not None:
        exec_rows.append((
            "Site Envelope (concept)",
            f"≈ {site_layout.envelope_w_m:.1f} × {site_layout.envelope_d_m:.1f} m",
        ))
    _add_table(doc, exec_rows, ["Metric", "Value"])
    # The generic average DC-to-AC split is meaningless for a governed run (each
    # governed block has a fixed DC count); §9 shows the governed composition.
    if ctx.dc_blocks_allocation and not ctx.configuration_code:
        alloc_parts = [
            f"{entry.get('dc_blocks_per_ac_block')} DC Blocks per AC Block × "
            f"{entry.get('ac_blocks_count')} AC Blocks"
            for entry in ctx.dc_blocks_allocation
        ]
        doc.add_paragraph(f"DC-to-AC allocation: {'; '.join(alloc_parts)}.")

    # --- Section 2: Project Inputs & Assumptions ---
    # Flows on the same page as Section 1 when space allows (no forced page break).
    # Note: POI power/energy/guarantee values are already in Section 1 — not repeated here.
    doc.add_heading("2.  Project Inputs & Assumptions", level=2)
    _keep_next_para(doc.add_paragraph(
        "Key design parameters assumed in this sizing. "
        "Efficiency values are one-way (DC → POI). Loss and DoD values exclude auxiliary loads."
    ))
    _keep_next_para(doc.add_paragraph(
        "The listed efficiency values do not include Auxiliary losses; no Auxiliary loss assumption is applied in this report."
    ))
    _dod = ctx.stage1.get("dod_frac") if isinstance(ctx.stage1, dict) else None
    _sc = ctx.stage1.get("sc_loss_frac") if isinstance(ctx.stage1, dict) else None
    _cyc = ctx.cycles_per_year or (ctx.stage1.get("cycles_per_year") if isinstance(ctx.stage1, dict) else None)
    _sc_months = ctx.stage1.get("sc_time_months") if isinstance(ctx.stage1, dict) else None
    site_rows = [
        ("Grid MV Voltage (kV)", format_value(ctx.grid_mv_voltage_kv_ac, "kV")),
        ("Grid Frequency (Hz)", _format_or_tbd(ctx.project_inputs.get("poi_frequency_hz"), "Hz")),
        ("Grid Power Factor", format_value(ctx.grid_power_factor, "PF")),
        ("Project Life (Years)", f"{ctx.project_life_years}"),
        ("Cycles per Year (Assumed)", f"{int(_cyc)}" if _cyc else "—"),
        ("Depth of Discharge (DoD)", format_percent(_dod, input_is_fraction=True) if _dod is not None else "—"),
        ("Shipping & Commissioning (S&C) Loss", format_percent(_sc, input_is_fraction=True) if _sc is not None else "—"),
    ]
    if _sc_months is not None:
        site_rows.append(("FAT-to-COD Duration", f"{int(round(_sc_months))} months"))
    _add_table(doc, site_rows, ["Parameter", "Value"])

    # --- Section 3: Stage 1 - DC Energy Sizing ---
    doc.add_page_break()
    doc.add_heading("3.  Stage 1 – DC Energy Sizing", level=2)
    _keep_next_para(doc.add_paragraph(
        "DC Energy Required (MWh) = POI Energy Requirement ÷ "
        "((1 − S&C loss) × DoD × DC RTE (discharge) × One-way Efficiency)"
    ))
    _keep_next_para(doc.add_paragraph(
        f"One-way Efficiency (DC→POI): {format_percent(ctx.efficiency_chain_oneway_frac, input_is_fraction=True)}  |  "
        f"S&C loss: {format_percent(ctx.stage1.get('sc_loss_frac') or 0.0, input_is_fraction=True)}  |  "
        f"DoD: {format_percent(ctx.stage1.get('dod_frac') or 0.0, input_is_fraction=True)}  |  "
        f"DC RTE: {format_percent(ctx.stage1.get('dc_round_trip_efficiency_frac') or 0.0, input_is_fraction=True)}"
    ))
    rte_adj = float(ctx.stage1.get("rte_curve_adjust_pp") or 0.0)
    if rte_adj != 0.0:
        _keep_next_para(doc.add_paragraph(f"RTE Curve Adjustment: {rte_adj:+.1f} pp"))
    s1_rows = [
        ("DC Energy Capacity Required (MWh)", format_value(ctx.stage1.get("dc_energy_capacity_required_mwh") or 0.0, "MWh")),
        ("DC Power Required (MW)", format_value(ctx.stage1.get("dc_power_required_mw") or 0.0, "MW")),
    ]
    _add_table(doc, s1_rows, ["Metric", "Value"])

    doc.add_heading("3.1  Efficiency Chain (One-way, DC → POI)", level=3)
    _keep_next_para(doc.add_paragraph(
        "Component efficiencies represent the one-way conversion path from DC terminals to POI. "
        "Their product equals the total one-way efficiency shown in the first row."
    ))
    eff_rows = [
        ("Total One-way Efficiency", format_percent(ctx.efficiency_chain_oneway_frac, input_is_fraction=True)),
        ("DC Cables", format_percent(ctx.efficiency_components_frac.get("eff_dc_cables_frac"), input_is_fraction=True)),
        ("PCS", format_percent(ctx.efficiency_components_frac.get("eff_pcs_frac"), input_is_fraction=True)),
        ("Transformer (MVT)", format_percent(ctx.efficiency_components_frac.get("eff_mvt_frac"), input_is_fraction=True)),
        ("RMU / Switchgear / AC Cables", format_percent(ctx.efficiency_components_frac.get("eff_ac_cables_sw_rmu_frac"), input_is_fraction=True)),
        ("HVT / Others", format_percent(ctx.efficiency_components_frac.get("eff_hvt_others_frac"), input_is_fraction=True)),
    ]
    _add_table(doc, eff_rows, ["Component", "Efficiency"])

    # --- Section 4: Stage 2 - DC Block Configuration ---
    # Flows after Stage 1 (no forced break): the section is one small table and
    # previously produced a nearly blank page of its own.
    doc.add_heading("4.  Stage 2 – DC Block Configuration", level=2)
    dc_table = ctx.stage2.get("block_config_table") if isinstance(ctx.stage2, dict) else None

    def _format_mwh_3(value):
        try:
            return f"{float(value):.3f}"
        except Exception:
            return "" if value is None else str(value)

    if dc_table is not None and not dc_table.empty:
        drop_cols = {"Config Adjustment (%)", "Oversize (MWh)", "Busbars Needed (K=10)", "Busbars Needed"}
        dc_columns = [c for c in dc_table.columns if c not in drop_cols]
        headers_map = {c: c for c in dc_columns}
        formatters = {}
        for col in dc_columns:
            if col in ("Unit Capacity (MWh)", "Subtotal (MWh)", "Total DC Nameplate @BOL (MWh)"):
                formatters[col] = _format_mwh_3
            elif col in ("Block Code", "Block Name"):
                formatters[col] = lambda v: neutralize_equipment_text(v, brand)
            else:
                formatters[col] = lambda v: "" if v is None else str(v)
        _add_dataframe_table(doc, dc_table, dc_columns, headers_map, formatters)
    else:
        doc.add_paragraph("DC block configuration table unavailable.")
    oversize_mwh = ctx.stage2.get("oversize_mwh") if isinstance(ctx.stage2, dict) else None
    doc.add_paragraph(
        f"DC total nameplate @BOL: {_format_mwh_3(ctx.dc_total_energy_mwh)} MWh   |   "
        f"Oversize margin: {_format_mwh_3(oversize_mwh)} MWh"
    )
    # State the DC packaging mode and, for a hybrid config, call out the cabinet
    # tail explicitly (the DC-side counterpart of the AC-side tail AC Blocks).
    _s2 = ctx.stage2 if isinstance(ctx.stage2, dict) else {}
    _dc_mode = str(_s2.get("mode") or "")
    _cont_n = int(_s2.get("container_count") or 0)
    _cab_n = int(_s2.get("cabinet_count") or 0)
    _mode_label = {
        "container_only": "Container only",
        "cabinet_only": "Cabinet only",
        "hybrid": "Hybrid (full containers + a cabinet tail)",
    }.get(_dc_mode, _dc_mode)
    if _dc_mode:
        _mode_txt = f"DC packaging: {_mode_label}"
        if _dc_mode == "hybrid":
            _mode_txt += f" — {_cont_n} container(s) + {_cab_n} cabinet(s) as the DC-side tail"
        elif _dc_mode == "container_only":
            _mode_txt += f" — {_cont_n} container(s)"
        elif _dc_mode == "cabinet_only":
            _mode_txt += f" — {_cab_n} cabinet(s)"
        doc.add_paragraph(_mode_txt + ".")

    # --- Section 5: Stage 3 - Lifetime Degradation & POI Deliverable ---
    doc.add_page_break()
    doc.add_heading("5.  Stage 3 – Lifetime Degradation & POI Deliverable", level=2)
    s3_meta = ctx.stage3_meta if isinstance(ctx.stage3_meta, dict) else {}
    if s3_meta:
        def _fmt_float(value, decimals=2, default=""):
            try:
                return f"{float(value):.{decimals}f}"
            except Exception:
                return default

        poi_power = s3_meta.get("poi_power_mw", ctx.poi_power_requirement_mw)
        dc_power = s3_meta.get("dc_power_mw")
        if dc_power is None and isinstance(ctx.stage1, dict):
            dc_power = ctx.stage1.get("dc_power_required_mw")
        eff_c_rate = s3_meta.get("effective_c_rate")
        chosen_soh_c_rate = s3_meta.get("chosen_soh_c_rate")
        chosen_cycles_per_year = s3_meta.get("chosen_soh_cycles_per_year")

        model_rows = [
            ("POI Power (MW)", _fmt_float(poi_power, 2)),
            ("DC-equivalent Power (MW)", _fmt_float(dc_power, 2)),
            ("Effective C-rate (DC side)", f"{_fmt_float(eff_c_rate, 3)} C"),
            ("SOH Modelling C-rate", f"≤ {chosen_soh_c_rate}" if chosen_soh_c_rate else "—"),
            ("SOH Modelling Cycles/Year", f"{chosen_cycles_per_year}" if chosen_cycles_per_year else "—"),
            ("Guarantee Year (from COD)", f"{ctx.poi_guarantee_year}"),
            ("POI Energy Guarantee Target (MWh)", format_value(gyr_target, "MWh")),
        ]
        _add_table(doc, model_rows, ["Parameter", "Value"])
        doc.add_paragraph("")

    s3_df = ctx.stage3_df
    if s3_df is not None and not s3_df.empty:
        rte_columns = ["DC_RTE_Pct", "System_RTE_Pct"]
        rte_frame = (
            s3_df.loc[:, rte_columns].apply(pd.to_numeric, errors="coerce").dropna()
            if set(rte_columns).issubset(s3_df.columns)
            else pd.DataFrame()
        )
        if not rte_frame.empty:
            doc.add_paragraph(
                "System RTE = DC RTE × (One-way Efficiency)², "
                "where one-way efficiency is the DC→POI chain defined in Section 3.1."
            )
            rte_dc = rte_frame["DC_RTE_Pct"]
            rte_sys = rte_frame["System_RTE_Pct"]
            rte_dc_min, rte_dc_max = float(rte_dc.min()), float(rte_dc.max())
            rte_sys_min, rte_sys_max = float(rte_sys.min()), float(rte_sys.max())

            if abs(rte_dc_min - rte_dc_max) < 1e-6:
                doc.add_paragraph(f"DC RTE: {format_percent(rte_dc_min, input_is_fraction=False)} (constant over project life)")
            else:
                doc.add_paragraph(
                    f"DC RTE: {format_percent(rte_dc_min, input_is_fraction=False)} "
                    f"to {format_percent(rte_dc_max, input_is_fraction=False)} over project life"
                )
            if abs(rte_sys_min - rte_sys_max) < 1e-6:
                doc.add_paragraph(f"System RTE: {format_percent(rte_sys_min, input_is_fraction=False)} (constant over project life)")
            else:
                doc.add_paragraph(
                    f"System RTE: {format_percent(rte_sys_min, input_is_fraction=False)} "
                    f"to {format_percent(rte_sys_max, input_is_fraction=False)} over project life"
                )
        else:
            doc.add_paragraph("DC and system RTE ranges are unavailable for the supplied Stage 3 dataset.")
        doc.add_paragraph("")

        s3_df = s3_df.copy()
        guarantee_year_int = int(ctx.poi_guarantee_year or 0)
        # The contractual guarantee is evaluated at the guarantee year only;
        # showing Yes/No for every year misreads as a 20-year failure list.
        s3_df["Meets_Guarantee"] = [
            ("Yes" if float(poi) >= float(gyr_target or 0.0) else "No")
            if int(year) == guarantee_year_int else "—"
            for year, poi in zip(s3_df["Year_Index"], s3_df["POI_Usable_Energy_MWh"])
        ]

        capacity_chart = _plot_dc_capacity_bar_png(
            bol_mwh=ctx.dc_total_energy_mwh,
            s3_df=s3_df,
            guarantee_year=int(ctx.poi_guarantee_year or 0),
            poi_target=gyr_target,
            title="Installed DC Energy vs. Deliverable POI Energy at Key Milestones",
        )
        if capacity_chart is not None and capacity_chart.getbuffer().nbytes > 0:
            doc.add_picture(capacity_chart, width=Inches(6.7))
            _keep_next_para(doc.paragraphs[-1])  # keep picture paragraph with its caption
            _keep_next_para(doc.add_paragraph(
                f"Figure 1: Installed DC nameplate energy (dark, DC side) versus deliverable "
                f"energy at the point of interconnection (light, POI side). The step from BOL "
                f"to POI reflects the one-way DC→POI efficiency chain, depth of discharge, "
                f"and S&C/calendar/cycle degradation; later POI bars additionally reflect ageing. "
                f"Red line = POI energy guarantee target ({format_value(gyr_target, 'MWh')} MWh), "
                f"applicable to the POI bars only."
            ))
            _keep_next_para(doc.add_paragraph(""))

        chart = _plot_poi_usable_png(
            s3_df,
            poi_target=gyr_target,
            title=f"POI Usable Energy vs. Year  (guarantee target = red line, Year {ctx.poi_guarantee_year})",
        )
        if chart is not None and chart.getbuffer().nbytes > 0:
            doc.add_picture(chart, width=Inches(6.7))
            _keep_next_para(doc.paragraphs[-1])  # keep picture paragraph with its caption
            _keep_next_para(doc.add_paragraph(
                f"Figure 2: POI Usable Energy (MWh) by year from COD. "
                f"Red line = guarantee target ({format_value(gyr_target, 'MWh')} MWh). "
                f"Guarantee evaluated at Year {ctx.poi_guarantee_year}."
            ))
        else:
            err = None
            try:
                err = ctx.stage3_meta.get("error") if isinstance(ctx.stage3_meta, dict) else None
            except Exception:
                err = None
            _keep_next_para(doc.add_paragraph(f"Stage 3 chart unavailable{': ' + err if err else '.'}"))

        _keep_next_para(doc.add_paragraph(""))
        _keep_next_para(doc.add_paragraph("Year-by-Year Degradation & POI Deliverable:"))
        s3_columns = [
            "Year_Index",
            "SOH_Display_Pct",
            "SOH_Absolute_Pct",
            "DC_Usable_MWh",
            "POI_Usable_Energy_MWh",
            "DC_RTE_Pct",
            "System_RTE_Pct",
            "Meets_Guarantee",
        ]
        s3_headers = {
            "Year_Index": "Year (From COD)",
            "SOH_Display_Pct": "SOH @ COD Baseline (%)",
            "SOH_Absolute_Pct": "SOH vs FAT (%)",
            "DC_Usable_MWh": "DC Usable (MWh)",
            "POI_Usable_Energy_MWh": "POI Usable (MWh)",
            "DC_RTE_Pct": "DC RTE (%)",
            "System_RTE_Pct": "System RTE (%)",
            "Meets_Guarantee": f"Meets Target (Y{int(ctx.poi_guarantee_year or 0)})",
        }
        s3_formatters = {
            "Year_Index": lambda v: f"{int(v)}",
            "SOH_Display_Pct": lambda v: format_percent(v, input_is_fraction=False),
            "SOH_Absolute_Pct": lambda v: format_percent(v, input_is_fraction=False),
            "DC_Usable_MWh": lambda v: format_value(v, "MWh"),
            "POI_Usable_Energy_MWh": lambda v: format_value(v, "MWh"),
            "DC_RTE_Pct": lambda v: format_percent(v, input_is_fraction=False),
            "System_RTE_Pct": lambda v: format_percent(v, input_is_fraction=False),
            "Meets_Guarantee": lambda v: "" if v is None else str(v),
        }
        s3_columns = [column for column in s3_columns if column in s3_df.columns]
        _add_dataframe_table(doc, s3_df, s3_columns, s3_headers, s3_formatters, keep_together=False)
        doc.add_paragraph(
            f"Note: the POI energy guarantee is contractually evaluated at Year "
            f"{int(ctx.poi_guarantee_year or 0)} only; later years are shown for reference "
            f"and carry no pass/fail meaning."
        )

    # --- Section 6: Stage 4 - AC Block Sizing ---
    # Flows after the Stage 3 lifetime table (no forced break) to avoid an
    # orphan page when the table ends near a page boundary.
    doc.add_heading("6.  Stage 4 – AC Block Sizing", level=2)
    ac_ratio = ctx.ac_output.get("selected_ratio") if isinstance(ctx.ac_output, dict) else None
    ac_pcs_kw = ctx.ac_output.get("pcs_kw") if isinstance(ctx.ac_output, dict) else None
    if ac_pcs_kw is None and isinstance(ctx.ac_output, dict):
        ac_pcs_kw = ctx.ac_output.get("pcs_power_kw")
    pcs_count_by_block = ctx.ac_output.get("pcs_count_by_block") if isinstance(ctx.ac_output, dict) else None
    # "Product-backed" = the transformer nameplate comes from a real catalogue
    # product (a governed configuration OR a catalogue product bound on the
    # auto-recommend trunk), so the report shows the nameplate rather than MW ÷ PF.
    is_governed = bool(ctx.configuration_code) or (
        isinstance(ctx.ac_output, dict) and (
            bool(ctx.ac_output.get("governed_configuration"))
            or bool(ctx.ac_output.get("ac_block_product_block_code"))
        )
    )
    governed_groups = ctx.ac_output.get("governed_groups") if isinstance(ctx.ac_output, dict) else None
    pcs_by_block_text = ""
    if isinstance(governed_groups, list) and governed_groups:
        # Mixed governed (Phase B): describe the true per-group breakdown, not the
        # uniform head that the SLD-authoritative output carries.
        pcs_by_block_text = "; ".join(
            f"{int(g.get('pcs_per_ac_block') or 0)} PCS × {int(g.get('ac_block_count') or 0)} block(s)"
            for g in governed_groups
        )
    elif isinstance(pcs_count_by_block, list) and pcs_count_by_block:
        distinct_counts = {int(v) for v in pcs_count_by_block}
        if len(distinct_counts) == 1:
            pcs_by_block_text = (
                f"{distinct_counts.pop()} per block, uniform across all "
                f"{len(pcs_count_by_block)} AC Blocks"
            )
        else:
            from collections import Counter
            grouped = Counter(int(v) for v in pcs_count_by_block)
            pcs_by_block_text = "; ".join(
                f"{count} PCS × {n_blocks} blocks" for count, n_blocks in sorted(grouped.items(), reverse=True)
            )

    # Transformer sizing basis. A governed run's nameplate is an owner-confirmed
    # product value (or an explicit TBD), NEVER an AC-power ÷ PF estimate — so the
    # generic "MW ÷ PF" basis must not appear on a governed report (it would
    # contradict the real product MVA, e.g. promote 10 MW / 0.9 to 11.11 MVA).
    if is_governed:
        if ctx.transformer_rating_kva:
            transformer_formula = (
                f"Manufacturer nameplate rating "
                f"{format_value(ctx.transformer_rating_kva / 1000.0, 'MVA')} MVA (from product datasheet)"
            )
        else:
            transformer_formula = "To be confirmed from the product datasheet"
        transformer_basis_label = "Transformer Basis"
    else:
        transformer_mva = None
        if ctx.grid_power_factor and ctx.grid_power_factor > 0 and ctx.ac_block_size_mw:
            transformer_mva = ctx.ac_block_size_mw / ctx.grid_power_factor
        transformer_formula = "n/a"
        if transformer_mva is not None and ctx.ac_block_size_mw and ctx.grid_power_factor:
            transformer_formula = (
                f"{format_value(ctx.ac_block_size_mw, 'MW')} MW ÷ {format_value(ctx.grid_power_factor, 'PF')} (PF)"
                f" = {format_value(transformer_mva, 'MVA')} MVA (estimate)"
            )
        transformer_basis_label = "Transformer Sizing Basis"

    transformer_rating_label = (
        "Transformer Rating (per lead AC Block; see Section 9 for the full schedule)"
        if (is_governed and isinstance(governed_groups, list) and len(governed_groups) > 1)
        else "Transformer Rating (per block)"
    )
    s4_rows = [
        ("AC:DC Ratio", ac_ratio or "—"),
        ("AC Block Template", _normalize_template_label(ctx.ac_block_template_id)),
        ("AC Block Size (MW)", format_value(ctx.ac_block_size_mw, "MW")),
        ("PCS per AC Block", f"{ctx.pcs_per_block:d}"),
    ]
    if pcs_by_block_text:
        s4_rows.append(("PCS Count by Block", pcs_by_block_text))
    if ac_pcs_kw:
        s4_rows.append(("PCS Unit Rating (kW)", f"{float(ac_pcs_kw):.0f} kW"))
    s4_rows += [
        ("Feeders per AC Block", f"{ctx.feeders_per_block:d}"),
        ("PCS AC Output Voltage (V_LL)", format_value(ctx.pcs_lv_voltage_v_ll_rms_ac, "V")),
        ("Total AC Blocks", f"{ctx.ac_blocks_total:d}"),
        ("Total PCS Modules", f"{ctx.pcs_modules_total:d}"),
        (transformer_rating_label, _format_transformer_rating(ctx.transformer_rating_kva)),
        (transformer_basis_label, transformer_formula),
    ]
    _add_table(doc, s4_rows, ["Parameter", "Value"])

    # Mixed AC Block station: a head AC Block model plus smaller tail model(s).
    # This is the AC-side counterpart of the DC container + cabinet tail, driven
    # off the manual per-AC-Block adjustment in AC Sizing. Uniform stations carry
    # a single-entry breakdown, so the schedule is shown only when it adds
    # information (more than one distinct AC Block model).
    ac_breakdown = (
        ctx.ac_output.get("ac_block_breakdown") if isinstance(ctx.ac_output, dict) else None
    )
    if (
        isinstance(ctx.ac_output, dict)
        and ctx.ac_output.get("ac_block_mixed")
        and isinstance(ac_breakdown, list)
        and len(ac_breakdown) > 1
    ):
        doc.add_heading("6.1  Mixed AC Block Station Schedule", level=3)
        _keep_next_para(doc.add_paragraph(
            "This station mixes AC Block models: a head model covers most DC Blocks "
            "and smaller tail model(s) cover the remainder — the AC-side counterpart "
            "of the DC container + cabinet tail in Section 4."
        ))
        _keep_next_para(doc.add_paragraph(
            "Concept / draft: a mixed station is a manual engineering adjustment. "
            "Per-model OEM product and transformer data are not individually confirmed — "
            "each model's transformer rating below is a MW ÷ PF estimate, not a bound "
            "product nameplate — and the SLD (§7) renders the representative Head AC Block "
            "fleet only. Confirm each model's product and transformer data before use as a "
            "formal engineering result."
        ))
        pf = ctx.grid_power_factor if (ctx.grid_power_factor and ctx.grid_power_factor > 0) else None
        schedule_rows = []
        for entry in ac_breakdown:
            block_mw = float(entry.get("block_mw") or 0.0)
            dc_each = entry.get("dc_blocks_each")
            if dc_each is None:
                dc_text = f"{entry.get('dc_blocks_min')}–{entry.get('dc_blocks_max')} per block"
            else:
                dc_text = f"{int(dc_each)} per block"
            xf = f"{block_mw / pf:.2f} MVA (est.)" if pf else "TBD"
            schedule_rows.append((
                str(entry.get("role") or "—"),
                f"{int(entry.get('count') or 0)}",
                f"{int(entry.get('pcs_count') or 0)} × {float(entry.get('pcs_kw') or 0):.0f} kW",
                f"{block_mw:.2f} MW",
                dc_text,
                xf,
            ))
        _add_table(
            doc,
            schedule_rows,
            ["Role", "Qty", "PCS", "AC Block Size", "DC Blocks", "Transformer (per block)"],
        )

    # --- Section 7: Single Line Diagram ---
    doc.add_page_break()
    doc.add_heading("7.  Single Line Diagram", level=2)
    _keep_next_para(doc.add_paragraph(
        "System electrical configuration: MV bus → RMU → Step-up Transformer → AC Block (DC blocks attached)."
    ))
    figure_index = 3  # Figures 1-2 used by Stage 3 milestone + POI charts above
    sld_embedded = False
    if ctx.sld_pro_png_bytes:
        _add_concept_figure(doc, ctx.sld_pro_png_bytes, width=Inches(6.7))
        _keep_next_para(doc.paragraphs[-1])
        doc.add_paragraph(f"Figure {figure_index}: Single Line Diagram – System Electrical Configuration — NOT FOR CONSTRUCTION")
        figure_index += 1
        sld_embedded = True
    elif ctx.sld_preview_svg_bytes:
        png_bytes = _svg_bytes_to_png(ctx.sld_preview_svg_bytes)
        if png_bytes:
            _add_concept_figure(doc, png_bytes, width=Inches(6.7))
            _keep_next_para(doc.paragraphs[-1])
            doc.add_paragraph(f"Figure {figure_index}: Single Line Diagram – System Electrical Configuration — NOT FOR CONSTRUCTION")
            figure_index += 1
            sld_embedded = True
    if not sld_embedded:
        doc.add_paragraph("SLD not generated. Please generate in the Single Line Diagram page.")

    # --- Section 8: Typical AC Block Arrangement ---
    # Primary figure is the rule-based L1 drawing (spacing from the
    # ArrangementRuleProfile). The legacy UI-generated artifact is embedded
    # only as a fallback when the rule-based render is unavailable, so the
    # report never shows two contradictory aisle dimensions.
    doc.add_page_break()
    doc.add_heading("8.  Typical AC Block Arrangement (Concept Only)", level=2)
    plan_png, plan_layout, dc_per_ac = None, None, 0
    is_bilateral = ctx.layout_variant == BILATERAL_LAYOUT_VARIANT
    if is_bilateral:
        # Governed configuration: route by layout_variant to the confirmed
        # bilateral 4+4 engine so the drawing matches the SLD topology and the
        # confirmed physical layout — never rebuilt from an average DC-per-AC.
        try:
            bilateral = compute_bilateral_layout(ctx.dc_blocks_total // max(1, ctx.ac_blocks_total))
            plan_svg = render_bilateral_plan_svg(bilateral, block_label=str(ctx.configuration_code or ""))
            plan_png = _svg_bytes_to_png(plan_svg.encode("utf-8"))
            plan_layout = bilateral
        except Exception:
            plan_png, plan_layout = None, None
    elif ctx.ac_blocks_total > 0 and ctx.dc_blocks_total > 0:
        # For a mixed station the fractional average matches no real AC Block,
        # so draw the Head AC Block (consistent with §6.1 and the SLD head
        # fleet); tail AC Block(s) are described in §6.1.
        dc_per_ac = _representative_dc_per_ac(ctx)
        try:
            plan_svg, plan_layout = render_ac_block_plan_svg(
                dc_per_ac, ARRANGEMENT_PROFILE
            )
            plan_png = _svg_bytes_to_png(plan_svg.encode("utf-8"))
        except Exception:
            plan_png, plan_layout = None, None

    if plan_png and plan_layout is not None and is_bilateral:
        _bilateral_model_label = (
            (ctx.ac_output or {}).get("ac_block_model_name")
            or ctx.configuration_code
            or "1:8 / 8 x 1250 kW PCS concept"
        )
        _keep_next_para(doc.add_paragraph(
            f"AC Block model {_bilateral_model_label}: central vertical "
            f"40 ft AC Block with west 4-DC and east 4-DC mirrored fields "
            f"(4 + 4 arrangement, one DC Block per PCS):"
        ))
        _add_concept_figure(doc, plan_png, width=Inches(6.7))
        _keep_next_para(doc.paragraphs[-1])
        _keep_next_para(doc.add_paragraph(
            f"Figure {figure_index}: Typical AC Block Arrangement — equipment envelope ≈ "
            f"{plan_layout.envelope_w_m:.2f} × {plan_layout.envelope_d_m:.2f} m. "
            f"Concept only; spacing and 40 ft dimensions provisional. — NOT FOR CONSTRUCTION"
        ))
        figure_index += 1
        for note in getattr(plan_layout, "provisional_notes", ()):
            doc.add_paragraph(f"Provisional: {note}.")
    elif plan_png and plan_layout is not None:
        if _mixed_head_entry(ctx) is not None:
            _keep_next_para(doc.add_paragraph(
                "Mixed AC Block station: the arrangement below is the representative "
                "Head AC Block. Tail AC Block model(s) differ and are listed in the "
                "Mixed AC Block Station Schedule (§6.1)."
            ))
        _keep_next_para(doc.add_paragraph(
            f"Rule-based typical arrangement ({dc_per_ac} × DC per block, "
            f"mirrored back-to-back pairs, doors facing outward aisles):"
        ))
        _add_concept_figure(doc, plan_png, width=Inches(6.7))
        _keep_next_para(doc.paragraphs[-1])
        _keep_next_para(doc.add_paragraph(
            f"Figure {figure_index}: Typical AC Block Arrangement (rule-based) — "
            f"envelope ≈ {plan_layout.envelope_w_m:.2f} × "
            f"{plan_layout.envelope_d_m:.2f} m. Concept only. — NOT FOR CONSTRUCTION"
        ))
        figure_index += 1
        basis_rows = [(item, f"{value} — {basis}")
                      for item, value, basis in ARRANGEMENT_PROFILE.basis]
        basis_rows.append((
            "Arrangement rule profile", ARRANGEMENT_PROFILE.market_label))
        _add_table(doc, basis_rows, ["Spacing parameter", "Value & code basis"])
    else:
        layout_png_bytes = ctx.layout_png_bytes
        if not layout_png_bytes and ctx.layout_svg_bytes:
            layout_png_bytes = _svg_bytes_to_png(ctx.layout_svg_bytes)
        if layout_png_bytes:
            _add_concept_figure(doc, layout_png_bytes, width=Inches(6.7))
            _keep_next_para(doc.paragraphs[-1])
            doc.add_paragraph(
                f"Figure {figure_index}: Typical AC Block Arrangement — Concept Only · NOT FOR CONSTRUCTION")
        else:
            doc.add_paragraph(
                "Typical AC Block Arrangement not generated. "
                "Generate it from the corresponding concept page.")

    # --- Section 9: Concept Site Layout + Provisional Equipment Schedule ---
    # For a governed run, draw the whole-site concept layout (every governed AC
    # Block at its real product footprint) and list the provisional equipment
    # schedule (per governed group, with the bound product + transformer). Values
    # come from the actual AC Sizing run on ctx — never re-derived here. A generic
    # run keeps the legacy L2 linear site array.
    if ctx.configuration_code:
        site_png = None
        run_like = None
        try:
            run_like = _governed_run_from_ctx(ctx)
            layout_svg = render_governed_site_layout_concept_svg(
                run_like, title=f"CONCEPT SITE LAYOUT · {ctx.configuration_code}"
            )
            site_png = _svg_bytes_to_png(layout_svg.encode("utf-8"))
        except Exception:
            site_png, run_like = None, None
        if site_png and run_like is not None:
            doc.add_page_break()
            doc.add_heading("9.  Concept Site Layout & Equipment Schedule (Provisional)", level=2)
            _keep_next_para(doc.add_paragraph(
                f"{ctx.dc_blocks_total} DC Blocks compose {run_like.ac_blocks_total} AC Block(s) "
                f"across {len(run_like.groups)} configuration(s), {run_like.ac_power_mw_total:.2f} MW "
                f"total. Blocks are placed at their actual product footprint; this is a concept "
                f"arrangement, not a construction site layout against a project boundary."
            ))
            _add_concept_figure(doc, site_png, width=Inches(6.7))
            _keep_next_para(doc.paragraphs[-1])
            _keep_next_para(doc.add_paragraph(
                f"Figure {figure_index}: Concept Site Layout (actual product footprints) — Concept Only · NOT FOR CONSTRUCTION"
            ))
            figure_index += 1

            _keep_next_para(doc.add_paragraph(
                "Equipment schedule. Values marked TBD await confirmation. Product status: "
                "“confirmed” = selected and confirmed; “provisional” = preliminary product match, "
                "to be confirmed; “—” = product to be selected."
            ))
            _confirm_label = {
                "owner_selected": "confirmed",
                "provisional_auto": "provisional",
                "none": "—",
            }
            # Head = the largest AC Block (most DC per block); the smaller blocks
            # covering the remainder are the AC-side tail (mirrors the DC cabinet
            # tail in §4). A single-group site has a head and no tail.
            _max_dc_per = max((int(g.dc_blocks_per_ac_block or 0) for g in run_like.groups), default=0)
            _is_mixed_ac = len(run_like.groups) > 1
            schedule_rows = []
            for g in run_like.groups:
                mva = (g.ac_output or {}).get("transformer_mva")
                vg = (g.ac_output or {}).get("transformer_vector_group")
                xf = f"{float(mva):g} MVA" if mva else "TBD"
                if vg:
                    xf += f" · {vg}"
                conf = _confirm_label.get(getattr(g, "product_confirmation", "none"), "—")
                product_cell = (g.bound_product_code or "TBD (no catalogue product)")
                if g.bound_product_code:
                    product_cell += f"  [{conf}]"
                role = "Head" if (int(g.dc_blocks_per_ac_block or 0) == _max_dc_per) else "Tail"
                if not _is_mixed_ac:
                    role = "—"
                schedule_rows.append((
                    g.configuration_code,
                    role,
                    f"{g.ac_block_count}",
                    f"{g.pcs_per_ac_block} × {g.pcs_kw:.0f} kW",
                    f"{g.dc_blocks_per_ac_block} per block",
                    xf,
                    product_cell,
                ))
            _add_table(
                doc,
                schedule_rows,
                ["AC Block Model", "Role", "Qty", "PCS", "DC Blocks", "Transformer (per block)", "Product"],
            )
            if _is_mixed_ac:
                _keep_next_para(doc.add_paragraph(
                    "Mixed AC Block station: a head AC Block plus smaller tail AC Blocks cover the "
                    "remainder — the AC-side counterpart of the DC cabinet tail in Section 4."
                ))
        # A governed run does not also draw the linear L2 site array.
        site_layout = None
        group_layout = None
    else:
        site_layout = _compute_site_layout(ctx)
        group_layout = _compute_typical_group_layout(ctx)
    if site_layout is not None and group_layout is not None:
        try:
            # Draw ONE representative project group at legible scale, not the
            # full-site strip (which fits a page only as an illegible sliver).
            site_svg = render_site_svg(group_layout, SITE_PROFILE)
            site_png = _svg_bytes_to_png(site_svg.encode("utf-8"))
        except Exception:
            site_png = None
        if site_png:
            # A mixed station has no single "average" block; the figure draws the
            # Head AC Block group and the whole-site power/energy come from the
            # actual head + tail sizing, not a uniform multiplication.
            _mixed_head = _mixed_head_entry(ctx)
            _ac_out = ctx.ac_output if isinstance(ctx.ac_output, dict) else {}
            ws_power_mw = site_layout.total_power_mw
            ws_energy_mwh = site_layout.total_energy_mwh
            ws_blocks_desc = f"{site_layout.n_blocks:d} × {site_layout.dc_per_block} DC/block"
            if _mixed_head is not None:
                if _ac_out.get("total_ac_mw"):
                    ws_power_mw = float(_ac_out["total_ac_mw"])
                if ctx.dc_total_energy_mwh:
                    ws_energy_mwh = float(ctx.dc_total_energy_mwh)
                ws_blocks_desc = f"{site_layout.n_blocks:d} (mixed head + tail — see §6.1)"
            doc.add_page_break()
            doc.add_heading("9.  Concept Site Arrangement (Concept Only)", level=2)
            if _mixed_head is not None:
                _keep_next_para(doc.add_paragraph(
                    "Mixed AC Block station: the representative group below is drawn "
                    "with the Head AC Block. Tail AC Block model(s) differ (see the "
                    "Mixed AC Block Station Schedule, §6.1); the whole-site power and "
                    "energy stated here reflect the full head + tail composition."
                ))
            _keep_next_para(doc.add_paragraph(
                f"The site repeats one project group, so the figure below shows a "
                f"single representative project group of {group_layout.n_blocks} × "
                f"AC Block ({group_layout.dc_per_block} × DC each) at legible scale. "
                f"The complete site comprises {site_layout.groups} such group(s) "
                f"({site_layout.n_blocks} blocks total); drawing every block at "
                f"page scale would be illegible. Within a group the two AC blocks in "
                f"each row are mirrored so both PCS & MV stations face a shared "
                f"central MV corridor; feeders collect there and route one direction "
                f"along the access road to the substation. Fire apparatus access "
                f"roads run between groups and along the perimeter — not between "
                f"every row — so each block stays within "
                f"{SITE_PROFILE.fire_access_limit_m:.0f} m of a road (worst case "
                f"{site_layout.fire_access_reach_m:.1f} m within a group)."
            ))
            # A single group is compact; cap by height only if it is deeper
            # than it is wide, otherwise fit to the text column width.
            if group_layout.envelope_d_m > group_layout.envelope_w_m:
                _add_concept_figure(doc, site_png, height=Inches(5.5))
            else:
                _add_concept_figure(doc, site_png, width=Inches(6.2))
            _keep_next_para(doc.paragraphs[-1])
            _keep_next_para(doc.add_paragraph(
                f"Figure {figure_index}: Concept Site Arrangement — one "
                f"representative project group (≈ {group_layout.envelope_w_m:.1f} × "
                f"{group_layout.envelope_d_m:.1f} m; {group_layout.n_blocks} × "
                f"{group_layout.total_power_mw / max(1, group_layout.n_blocks):.0f} MW "
                f"per block). The full site is {site_layout.groups} such group(s) "
                f"≈ {site_layout.envelope_w_m:.1f} × {site_layout.envelope_d_m:.1f} m "
                f"({ws_power_mw:.0f} MW / "
                f"{ws_energy_mwh:.1f} MWh). Concept only — a full "
                f"Master Layout requires a registered site constraint set. — NOT FOR CONSTRUCTION"
            ))
            figure_index += 1
            site_rows = [
                ("AC Blocks (whole site)", ws_blocks_desc),
                ("Project groups", f"{site_layout.groups} group(s) × ≤ "
                             f"{site_layout.blocks_per_group} blocks · "
                             f"{site_layout.fire_roads} internal fire road(s)"),
                ("Representative group (drawn)",
                 f"{group_layout.n_blocks} block(s) · ≈ "
                 f"{group_layout.envelope_w_m:.1f} × {group_layout.envelope_d_m:.1f} m"),
                ("Fire access reach", f"≤ {site_layout.fire_access_reach_m:.1f} m "
                                      f"(limit {SITE_PROFILE.fire_access_limit_m:.0f} m)"),
                ("Site envelope (concept, whole site)",
                 f"≈ {site_layout.envelope_w_m:.1f} × {site_layout.envelope_d_m:.1f} m"),
                ("Rated power / energy (whole site)",
                 f"{ws_power_mw:.0f} MW / "
                 f"{ws_energy_mwh:.1f} MWh"),
            ]
            site_rows.extend(
                (item, f"{value} — {basis}")
                for item, value, basis in SITE_PROFILE.basis
            )
            _add_table(doc, site_rows, ["Site parameter", "Value & code basis"])

    return _doc_to_bytes(doc)


export_report_v2 = export_report_v2_1
