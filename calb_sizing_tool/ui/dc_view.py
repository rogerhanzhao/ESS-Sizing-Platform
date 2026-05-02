# -----------------------------------------------------------------------------
# Personal Open-Source Notice
#
# Copyright (c) 2026 Alex.Zhao. All rights reserved.
#
# This repository is released under the MIT License (see LICENSE file).
# Intended use: learning, evaluation, and engineering reference for Utility-scale
# BESS/ESS sizing and Reporting workflows.
#
# DISCLAIMER: This software is provided "AS IS", without warranty of any kind,
# express or implied. In no event shall the author(s) be liable for any claim,
# damages, or other liability arising from, out of, or in connection with the
# software or the use or other dealings in the software.
#
# NOTE: This is a personal project. It is not an official product or statement
# of any company or organization.
# -----------------------------------------------------------------------------

import streamlit as st
import pandas as pd
import numpy as np
import altair as alt
import os
import io
from pathlib import Path

# --- Adapted imports for refactor ---
from calb_sizing_tool.adapters.excel_loader_adapter import bundle_to_legacy_tuple, load_dc_excel_bundle_from_path
from calb_sizing_tool.config import DC_DATA_PATH, DC_DATA_IS_LEGACY, PROJECT_ROOT
from calb_sizing_tool.common.nameplate import apply_block_nameplate_recalc, get_standard_container_mwh
from calb_sizing_tool.infra.db.session import session_scope
from calb_sizing_tool.schemas.case import SizingCaseInput
from calb_sizing_tool.schemas.master_data import DcExcelMasterDataBundle
from calb_sizing_tool.schemas.run_snapshot import DcPipelineRunSnapshot
from calb_sizing_tool.schemas.stage1 import Stage1Result
from calb_sizing_tool.schemas.stage2 import Stage2Result
from calb_sizing_tool.schemas.stage3 import Stage3Result
from calb_sizing_tool.services.access_control_service import AccessControlService
from calb_sizing_tool.services.dc_pipeline_service import size_with_guarantee as service_size_with_guarantee
from calb_sizing_tool.services.run_persistence_service import persist_dc_run
from calb_sizing_tool.services.stage1_service import run_stage1 as service_run_stage1
from calb_sizing_tool.services.stage2_service import (
    K_MAX_FIXED,
    build_config_cabinet_only as service_build_config_cabinet_only,
    build_config_container_only as service_build_config_container_only,
    build_config_hybrid as service_build_config_hybrid,
    pick_dc_block as service_pick_dc_block,
)
from calb_sizing_tool.services.stage3_service import (
    run_stage3 as service_run_stage3,
    select_rte_profile as service_select_rte_profile,
    select_soh_profile as service_select_soh_profile,
)
from calb_sizing_tool.ui.stage4_interface import pack_stage13_output
# Model import for AC/SLD handoff
from calb_sizing_tool.models import DCBlockResult
from calb_sizing_tool.state.project_state import bump_run_id_dc, init_project_state
from calb_sizing_tool.state.auth_state import get_auth_context, get_auth_user
from calb_sizing_tool.state.session_state import init_shared_state, set_run_time
from calb_sizing_tool.state.workspace_state import get_workspace_context, restore_run_bundle_to_session

# ==========================================
# 0. SETUP & LIBRARY CHECK
# ==========================================
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Optional: Matplotlib for DOCX chart export
try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except Exception:
    MATPLOTLIB_AVAILABLE = False

# ==========================================
# CALB VI COLORS
# ==========================================
CALB_SKY_BLUE   = "#5cc3e4"
CALB_DEEP_BLUE  = "#23496b"
CALB_BLACK      = "#1e1e1e"
CALB_GREY       = "#58595b"
CALB_LIGHT_GREY = "#bebec3"
CALB_WISE_GREY  = "#cedbea"
CALB_WHITE      = "#ffffff"

# ==========================================
# HELPERS
# ==========================================
def inject_css():
    base_theme = st.get_option("theme.base") or "light"
    is_dark = (base_theme == "dark")

    if is_dark:
        BG_MAIN      = "#0f1116"
        TITLE_COLOR  = CALB_SKY_BLUE
        CARD_BG      = "#1b1e24"
        CARD_BORDER  = "#2c2f36"
        METRIC_LABEL = "#9ca3af"
        METRIC_VALUE = CALB_SKY_BLUE
        BUTTON_BG    = CALB_SKY_BLUE
        BUTTON_FG    = CALB_BLACK
        BUTTON_HOVER_BG = CALB_WHITE
        BUTTON_HOVER_FG = CALB_DEEP_BLUE
    else:
        BG_MAIN      = CALB_WISE_GREY
        TITLE_COLOR  = CALB_DEEP_BLUE
        CARD_BG      = CALB_WHITE
        CARD_BORDER  = CALB_LIGHT_GREY
        METRIC_LABEL = CALB_GREY
        METRIC_VALUE = CALB_DEEP_BLUE
        BUTTON_BG    = CALB_SKY_BLUE
        BUTTON_FG    = CALB_WHITE
        BUTTON_HOVER_BG = CALB_DEEP_BLUE
        BUTTON_HOVER_FG = CALB_WHITE

    global CHART_TEXT_COLOR
    CHART_TEXT_COLOR = CALB_DEEP_BLUE if not is_dark else CALB_SKY_BLUE

    st.markdown(
        f"""
    <style>
    .calb-page-title {{
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
        color: {TITLE_COLOR};
        margin-bottom: 0.0rem;
    }}
    .calb-card {{
        background-color: {CARD_BG};
        padding: 1.2rem 1.6rem;
        border-radius: 18px;
        border: 1px solid {CARD_BORDER};
        box-shadow: 0 6px 18px rgba(0, 0, 0, 0.15);
        margin-bottom: 1.4rem;
    }}
    .metric-label {{
        color: {METRIC_LABEL} !important;
        font-size: 0.85rem;
        text-transform: uppercase;
        letter-spacing: 0.04em;
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    }}
    .metric-value {{
        font-size: 1.3rem;
        font-weight: 600;
        color: {METRIC_VALUE} !important;
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    }}
    .stButton>button {{
        background-color: {BUTTON_BG};
        color: {BUTTON_FG};
        border-radius: 999px;
        border: none;
        padding: 0.7rem 2.3rem;
        font-weight: 650;
        font-size: 1.0rem;
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
        box-shadow: 0 4px 10px rgba(0,0,0,0.25);
    }}
    .stButton>button:hover {{
        background-color: {BUTTON_HOVER_BG};
        color: {BUTTON_HOVER_FG};
    }}
    </style>
    """,
        unsafe_allow_html=True,
    )

def to_float(x, default=0.0):
    try:
        if isinstance(x, str):
            x = x.replace("%", "").replace(",", "").strip()
        return float(x)
    except Exception:
        return default

def to_int(x, default=0):
    try:
        return int(to_float(x, default))
    except Exception:
        return default

def to_frac(x, default=1.0):
    v = to_float(x, default)
    if v > 1.5:
        return v / 100.0
    return v

def clamp01(value: float) -> float:
    try:
        return min(1.0, max(0.0, float(value)))
    except Exception:
        return 0.0

def safe_div(a: float, b: float, default: float = 0.0) -> float:
    try:
        if b is None or abs(float(b)) < 1e-12:
            return default
        return float(a) / float(b)
    except Exception:
        return default

def calc_sc_loss_pct(sc_months: float) -> float:
    m = to_int(sc_months, 0)
    if m <= 0:
        return 0.0

    mapping = {
        1: 2.0, 2: 2.0, 3: 2.0,
        4: 2.5, 5: 2.8, 6: 3.0,
        7: 3.2, 8: 3.5, 9: 3.8,
        10: 4.1, 11: 4.3, 12: 4.5,
    }

    if m in mapping:
        return mapping[m]

    if m > 12:
        base_loss = 4.5
        extra_months = m - 12
        return base_loss + (extra_months * 0.05)

    return mapping.get(m, 2.0)

def first_success_key(results: dict, preferred_order: list):
    for k in preferred_order:
        v = results.get(k)
        if isinstance(v, tuple) and len(v) > 0 and v[0] != "ERROR":
            return k
    return None

# ==========================================
# 3. LOAD DATA FROM EXCEL
# ==========================================
@st.cache_data
def load_data(path: Path):
    bundle = load_dc_excel_bundle_from_path(path)
    return bundle_to_legacy_tuple(bundle)

def validate_rte_monotonicity_314(
    df_rte_profile: pd.DataFrame,
    df_rte_curve: pd.DataFrame,
    tol: float = 0.0005,
) -> pd.DataFrame:
    required_profile = {"Profile_Id", "C_Rate"}
    required_curve = {"Profile_Id", "Soh_Band_Min_Pct", "Rte_Dc_Pct"}
    if not required_profile.issubset(df_rte_profile.columns):
        return pd.DataFrame(columns=["Soh_Band_Min_Pct", "C_Rate", "Profile_Id", "Rte_Dc_Pct", "Expected_Max_Rte"])
    if not required_curve.issubset(df_rte_curve.columns):
        return pd.DataFrame(columns=["Soh_Band_Min_Pct", "C_Rate", "Profile_Id", "Rte_Dc_Pct", "Expected_Max_Rte"])

    prof = df_rte_profile[list(required_profile)].copy()
    prof["C_Rate"] = prof["C_Rate"].apply(lambda x: to_float(x, np.nan))

    curve = df_rte_curve[list(required_curve)].copy()
    curve["Soh_Band_Min_Pct"] = curve["Soh_Band_Min_Pct"].apply(lambda x: to_frac(x))
    curve["Rte_Dc_Pct"] = curve["Rte_Dc_Pct"].apply(lambda x: to_frac(x))

    merged = curve.merge(prof, on="Profile_Id", how="left")
    merged = merged.dropna(subset=["C_Rate", "Soh_Band_Min_Pct", "Rte_Dc_Pct"])
    if merged.empty:
        return pd.DataFrame(columns=["Soh_Band_Min_Pct", "C_Rate", "Profile_Id", "Rte_Dc_Pct", "Expected_Max_Rte"])

    issues = []
    for soh_band, group in merged.groupby("Soh_Band_Min_Pct", dropna=True):
        group_sorted = group.sort_values("C_Rate")
        max_prev = None
        for _, row in group_sorted.iterrows():
            rte = float(row["Rte_Dc_Pct"])
            if max_prev is None:
                max_prev = rte
                continue
            expected_max = max_prev
            if rte > expected_max + tol:
                issues.append(
                    {
                        "Soh_Band_Min_Pct": float(soh_band) * 100.0,
                        "C_Rate": float(row["C_Rate"]),
                        "Profile_Id": int(row["Profile_Id"]),
                        "Rte_Dc_Pct": rte * 100.0,
                        "Expected_Max_Rte": expected_max * 100.0,
                    }
                )
            if rte > max_prev:
                max_prev = rte

    return pd.DataFrame(
        issues,
        columns=["Soh_Band_Min_Pct", "C_Rate", "Profile_Id", "Rte_Dc_Pct", "Expected_Max_Rte"],
    )

# ==========================================
# 4. CORE CALC LOGIC
# ==========================================
def run_stage1(inputs: dict, defaults: dict) -> dict:
    return service_run_stage1(inputs, defaults).to_legacy_dict()

def _pick_dc_block(df_blocks: pd.DataFrame, form: str):
    return service_pick_dc_block(df_blocks, form)

def build_config_container_only(required_dc_mwh: float, container_unit: float, container_code: str, container_name: str):
    return service_build_config_container_only(required_dc_mwh, container_unit, container_code, container_name).to_legacy_dict()

def build_config_cabinet_only(required_dc_mwh: float, cab_unit: float, cab_code: str, cab_name: str, k_max: int = K_MAX_FIXED):
    return service_build_config_cabinet_only(required_dc_mwh, cab_unit, cab_code, cab_name, k_max=k_max).to_legacy_dict()

def build_config_hybrid(required_dc_mwh: float,
                        container_unit: float, container_code: str, container_name: str,
                        cab_unit: float, cab_code: str, cab_name: str,
                        k_max: int = K_MAX_FIXED):
    return service_build_config_hybrid(
        required_dc_mwh,
        container_unit,
        container_code,
        container_name,
        cab_unit,
        cab_code,
        cab_name,
        k_max=k_max,
    ).to_legacy_dict()

def select_soh_profile(effective_c_rate: float, cycles_per_year: int, df_soh_profile: pd.DataFrame):
    return service_select_soh_profile(effective_c_rate, cycles_per_year, df_soh_profile)

def select_rte_profile(effective_c_rate: float, df_rte_profile: pd.DataFrame):
    return service_select_rte_profile(effective_c_rate, df_rte_profile)

def run_stage3(stage1: dict,
               stage2: dict,
               df_soh_profile: pd.DataFrame,
               df_soh_curve: pd.DataFrame,
               df_rte_profile: pd.DataFrame,
               df_rte_curve: pd.DataFrame):
    bundle = DcExcelMasterDataBundle(
        workbook_path=Path(DC_DATA_PATH),
        defaults={},
        df_blocks=pd.DataFrame(),
        df_soh_profile=df_soh_profile.copy(),
        df_soh_curve=df_soh_curve.copy(),
        df_rte_profile=df_rte_profile.copy(),
        df_rte_curve=df_rte_curve.copy(),
        raw_sheets={},
    )
    stage1_model = Stage1Result.model_validate(stage1)
    records = stage2.get("block_config_table_records")
    if records is None and isinstance(stage2.get("block_config_table"), pd.DataFrame):
        records = stage2["block_config_table"].to_dict("records")
    stage2_payload = {
        "mode": stage2.get("mode"),
        "dc_nameplate_bol_mwh": stage2.get("dc_nameplate_bol_mwh"),
        "oversize_mwh": stage2.get("oversize_mwh"),
        "config_adjustment_frac": stage2.get("config_adjustment_frac"),
        "container_count": stage2.get("container_count"),
        "cabinet_count": stage2.get("cabinet_count"),
        "busbars_needed": stage2.get("busbars_needed"),
        "block_config_items": [
            {
                "block_code": record.get("Block Code"),
                "block_name": record.get("Block Name"),
                "form": record.get("Form"),
                "unit_capacity_mwh": record.get("Unit Capacity (MWh)"),
                "count": record.get("Count"),
                "subtotal_mwh": record.get("Subtotal (MWh)", 0.0),
                "total_dc_nameplate_bol_mwh": record.get("Total DC Nameplate @BOL (MWh)", stage2.get("dc_nameplate_bol_mwh", 0.0)),
            }
            for record in (records or [])
        ],
    }
    stage2_model = Stage2Result.model_validate(stage2_payload)
    stage3_result = service_run_stage3(stage1_model, stage2_model, bundle)
    return stage3_result.dataframe(), stage3_result.meta.to_legacy_dict()

def size_with_guarantee(stage1: dict,
                        mode: str,
                        df_blocks: pd.DataFrame,
                        df_soh_profile: pd.DataFrame,
                        df_soh_curve: pd.DataFrame,
                        df_rte_profile: pd.DataFrame,
                        df_rte_curve: pd.DataFrame,
                        k_max: int = K_MAX_FIXED,
                        max_iter: int = 60):
    bundle = DcExcelMasterDataBundle(
        workbook_path=Path(DC_DATA_PATH),
        defaults={},
        df_blocks=df_blocks.copy(),
        df_soh_profile=df_soh_profile.copy(),
        df_soh_curve=df_soh_curve.copy(),
        df_rte_profile=df_rte_profile.copy(),
        df_rte_curve=df_rte_curve.copy(),
        raw_sheets={},
    )
    snapshot = service_size_with_guarantee(
        Stage1Result.model_validate(stage1),
        mode,
        bundle,
        k_max=k_max,
        max_iter=max_iter,
    )
    s2 = snapshot.stage2.to_legacy_dict()
    s3_df = snapshot.stage3.dataframe()
    s3_meta = snapshot.stage3.meta.to_legacy_dict()
    return (
        s2,
        s3_df,
        s3_meta,
        snapshot.iteration_count,
        snapshot.poi_usable_energy_mwh_at_guarantee_year,
        snapshot.converged,
    )

# ==========================================
# 5. REPORT EXPORT HELPERS
# ==========================================
def find_logo_for_report():
    try:
        candidates = [
            PROJECT_ROOT / "calb_assets" / "logo" / "calb_logo.png",
            PROJECT_ROOT / "calb_logo.png",
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)
        data_dir = os.path.dirname(os.path.abspath(DC_DATA_PATH))
        for fname in os.listdir(data_dir):
            lower = fname.lower()
            if lower.endswith((".png", ".jpg", ".jpeg")) and ("logo" in lower or "calb" in lower):
                return os.path.join(data_dir, fname)
    except Exception:
        return None
    return None

def make_report_filename(project_name: str) -> str:
    base = (project_name or "CALB_ESS_Project").strip()
    if not base:
        base = "CALB_ESS_Project"
    safe = "".join(c if c.isalnum() or c in (" ", "_", "-") else "_" for c in base)
    safe = "_".join(safe.split())
    return f"{safe}_ESS_Sizing_Report.docx"

def _docx_add_config_table(doc: Document, df_conf: pd.DataFrame):
    if df_conf is None or df_conf.empty:
        doc.add_paragraph("No DC block configuration selected.", style="Intense Quote")
        return

    drop_cols = [
        "Block Code",
        "Form",
        "Config Adjustment (%)",
        "Busbars Needed (K=10)",
        "Oversize (MWh)",
        "Oversize_mwh",
        "Busbars Needed",
    ]
    df_show = df_conf.copy()
    cols_to_drop = [c for c in drop_cols if c in df_show.columns]
    if cols_to_drop:
        df_show = df_show.drop(columns=cols_to_drop)

    TOTAL_COL = "Total DC Nameplate @BOL (MWh)"
    has_total_col = TOTAL_COL in df_show.columns

    tbl = doc.add_table(rows=1, cols=len(df_show.columns))
    tbl.style = "Table Grid"

    hdr = tbl.rows[0].cells
    for j, col in enumerate(df_show.columns):
        hdr[j].text = str(col)
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for i, (_, row) in enumerate(df_show.iterrows()):
        rc = tbl.add_row().cells
        for j, col in enumerate(df_show.columns):
            if has_total_col and col == TOTAL_COL and i > 0:
                rc[j].text = ""
                continue

            val = row[col]
            if isinstance(val, float):
                rc[j].text = f"{val:.3f}"
            else:
                rc[j].text = f"{val}"

def _plot_poi_usable_png(s3_df: pd.DataFrame, poi_target: float, title: str) -> io.BytesIO:
    df = s3_df.sort_values("Year_Index").copy()
    x = df["Year_Index"].astype(int).tolist()
    y = df["POI_Usable_Energy_MWh"].astype(float).tolist()

    fig = plt.figure(figsize=(7.0, 3.2))
    ax = fig.add_subplot(111)
    ax.bar(x, y, color=CALB_SKY_BLUE)
    ax.axhline(poi_target, linewidth=2, color="#ff0000")
    ax.set_title(title)
    ax.set_xlabel("Year (from COD)")
    ax.set_ylabel("POI Usable Energy (MWh)")
    ax.set_xticks(x)
    ax.grid(True, axis="y", linestyle="--", alpha=0.35)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def _plot_dc_capacity_bar_png(
    s2: dict, s3_df: pd.DataFrame, guarantee_year: int, title: str
) -> io.BytesIO | None:
    if not MATPLOTLIB_AVAILABLE:
        return None
    try:
        bol = None
        if isinstance(s2, dict):
            bol = s2.get("dc_nameplate_bol_mwh")
        cod = None
        yx = None
        if s3_df is not None and not s3_df.empty:
            year0 = s3_df[s3_df["Year_Index"] == 0]
            if not year0.empty:
                cod = float(year0["POI_Usable_Energy_MWh"].iloc[0])
            g_row = s3_df[s3_df["Year_Index"] == int(guarantee_year)]
            if not g_row.empty:
                yx = float(g_row["POI_Usable_Energy_MWh"].iloc[0])

        labels = ["BOL", "COD", f"Y{int(guarantee_year)}"]
        values = [
            float(bol) if bol is not None else 0.0,
            float(cod) if cod is not None else 0.0,
            float(yx) if yx is not None else 0.0,
        ]

        fig = plt.figure(figsize=(6.6, 3.0))
        ax = fig.add_subplot(111)
        ax.bar(labels, values, color=CALB_SKY_BLUE)
        ax.set_title(title)
        ax.set_xlabel("Stage")
        ax.set_ylabel("Energy (MWh)")
        ax.grid(True, axis="y", linestyle="--", alpha=0.35)

        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=150)
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None

def _docx_add_lifetime_table(doc: Document, s3_df: pd.DataFrame):
    cols = [
        "Year_Index",
        "SOH_Display_Pct",
        "SOH_Absolute_Pct",
        "DC_Usable_MWh",
        "POI_Usable_Energy_MWh",
        "DC_RTE_Pct",
        "System_RTE_Pct",
    ]
    for c in cols:
        if c not in s3_df.columns:
            s3_df[c] = np.nan

    headers_map = {
        "Year_Index": "Year (From COD)",
        "SOH_Display_Pct": "SOH @ COD Baseline (%)",
        "SOH_Absolute_Pct": "SOH Vs FAT (%)",
        "DC_Usable_MWh": "DC Usable (MWh)",
        "POI_Usable_Energy_MWh": "POI Usable (MWh)",
        "DC_RTE_Pct": "DC RTE (%)",
        "System_RTE_Pct": "System RTE (%)",
    }

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for j, c in enumerate(cols):
        hdr[j].text = headers_map[c]
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    df_sorted = s3_df.sort_values("Year_Index")
    for _, r in df_sorted.iterrows():
        rc = tbl.add_row().cells
        rc[0].text = str(int(r["Year_Index"]))
        rc[1].text = f"{r['SOH_Display_Pct']:.1f}" if not pd.isna(r["SOH_Display_Pct"]) else ""
        rc[2].text = f"{r['SOH_Absolute_Pct']:.1f}" if not pd.isna(r["SOH_Absolute_Pct"]) else ""
        rc[3].text = f"{r['DC_Usable_MWh']:.2f}" if not pd.isna(r["DC_Usable_MWh"]) else ""
        rc[4].text = f"{r['POI_Usable_Energy_MWh']:.2f}" if not pd.isna(r["POI_Usable_Energy_MWh"]) else ""
        rc[5].text = f"{r['DC_RTE_Pct']:.2f}%" if not pd.isna(r["DC_RTE_Pct"]) else ""
        rc[6].text = f"{r['System_RTE_Pct']:.2f}%" if not pd.isna(r["System_RTE_Pct"]) else ""

def build_report_bytes(stage1: dict, results_dict: dict, report_order: list):
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    logo_path = find_logo_for_report()
    header = section.header
    header.is_linked_to_previous = False
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    hdr_cells = header_table.rows[0].cells

    if logo_path:
        p_logo = hdr_cells[0].paragraphs[0]
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.2))

    p_info = hdr_cells[1].paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_info = p_info.add_run(
        "CALB Group Co., Ltd.\n"
        "Utility-Scale Energy Storage Systems\n"
        "Confidential Sizing Report"
    )
    run_info.font.size = Pt(9)
    run_info.font.bold = True

    project_name = stage1.get("project_name", "CALB ESS Project")

    doc.add_heading("CALB Utility-Scale ESS Sizing Report", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(
        f"POI Requirement: {stage1['poi_power_req_mw']:.2f} MW / "
        f"{stage1['poi_energy_req_mwh']:.2f} MWh"
    )
    doc.add_paragraph("")

    doc.add_heading("1. Project Summary", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Project life: {int(stage1['project_life_years'])} years\n")
    p.add_run(f"POI guarantee year: {int(stage1.get('poi_guarantee_year', 0))}\n")
    p.add_run(f"Cycles per year (assumed): {int(stage1['cycles_per_year'])}\n")
    p.add_run(f"S&C time from FAT to COD: {int(round(stage1.get('sc_time_months', 0)))} months\n")
    p.add_run(f"RTE Curve Adjustment (pp): {float(stage1.get('rte_curve_adjust_pp', 0.0)):.1f}\n")
    p.add_run(f"DC->POI efficiency chain (one-way): {stage1.get('eff_dc_to_poi_frac', 0.0)*100:.2f}%\n")
    p.add_run(f"POI->DC equivalent power: {stage1.get('dc_power_required_mw', 0.0):.2f} MW")

    doc.add_paragraph(
        "This sizing report is based on the 314 Ah cell database and the internal "
        "CALB SOH/RTE profiles for the selected operating conditions."
    )

    doc.add_heading("2. Equipment Summary (DC Blocks)", level=2)

    for key, title in report_order:
        if key not in results_dict:
            continue
        s2, _, _, iter_count, poi_g, converged = results_dict[key]
        doc.add_paragraph(title, style=None)
        _docx_add_config_table(doc, s2.get("block_config_table"))
        doc.add_paragraph(f"Iterations: {iter_count} | Guarantee met: {bool(converged)}")
        if poi_g is not None:
            doc.add_paragraph(f"POI usable energy @ guarantee year: {poi_g:.2f} MWh")
        doc.add_paragraph("")

    doc.add_heading("3. Lifetime POI Usable Energy & SOH (Per Configuration)", level=2)

    poi_target = float(stage1["poi_energy_req_mwh"])
    guarantee_year = int(stage1.get("poi_guarantee_year", 0))

    for key, title in report_order:
        if key not in results_dict:
            continue
        _, s3_df, s3_meta, _, _, _ = results_dict[key]

        doc.add_paragraph(title, style=None)
        doc.add_paragraph(
            f"POI Power = {s3_meta.get('poi_power_mw', 0.0):.2f} MW | "
            f"DC-equivalent Power = {s3_meta.get('dc_power_mw', 0.0):.2f} MW | "
            f"Effective C-rate (DC-side) = {s3_meta.get('effective_c_rate', 0.0):.3f} C"
        )
        doc.add_paragraph(
            f"SOH profile ID = {s3_meta.get('soh_profile_id')} "
            f"(C-rate <= {s3_meta.get('chosen_soh_c_rate')}, cycles/year = {s3_meta.get('chosen_soh_cycles_per_year')}); "
            f"RTE profile ID = {s3_meta.get('rte_profile_id')} (C-rate <= {s3_meta.get('chosen_rte_c_rate')})."
        )
        doc.add_paragraph(f"Guarantee Year (from COD) = {guarantee_year} | POI Energy Target = {poi_target:.2f} MWh")

        if MATPLOTLIB_AVAILABLE:
            try:
                png = _plot_poi_usable_png(
                    s3_df=s3_df,
                    poi_target=poi_target,
                    title=f"POI Usable Energy vs Year - {key}"
                )
                doc.add_picture(png, width=Inches(6.7))
            except Exception:
                doc.add_paragraph("Chart export skipped due to plotting error.")
        else:
            doc.add_paragraph("Chart export skipped (matplotlib not available).")

        _docx_add_lifetime_table(doc, s3_df)
        doc.add_paragraph("")

    p_final = doc.add_paragraph()
    p_fmt = p_final.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    p_fmt.line_spacing = Pt(0)
    run_final = p_final.add_run()
    run_final.font.size = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ==========================================
# 6. MAIN VIEW FUNCTION
# ==========================================
def show():
    state = init_shared_state()
    init_project_state()
    dc_inputs = state.dc_inputs
    dc_results = state.dc_results
    ac_inputs = state.ac_inputs

    auth_context = get_auth_context()
    if auth_context is None:
        st.error("Login required.")
        return
    auth_user = get_auth_user()

    workspace = get_workspace_context()
    active_project_id = workspace.get("project_id")
    active_project_name = workspace.get("project_name")
    active_project_code = workspace.get("project_code")
    active_case_id = workspace.get("case_id")
    active_case_name = workspace.get("case_name")
    active_case_code = workspace.get("case_code")

    if active_project_id:
        with session_scope() as session:
            access = AccessControlService(session, auth_user)
            try:
                access.ensure_project_access(active_project_id)
            except PermissionError:
                st.error("You do not have access to the active project.")
                return

    if not active_project_id or not active_case_id:
        st.warning("Select a project and case before running DC sizing.")
        return

    st.caption(f"Active Project: {active_project_name} ({active_project_code})")
    st.caption(f"Active Case: {active_case_name} ({active_case_code})")

    # One-time default migration for UI fields (no impact on sizing logic)
    def _migrate_dc_defaults() -> None:
        version_key = "dc_defaults_version"
        current_version = int(st.session_state.get(version_key, 0) or 0)
        target_version = 4
        if current_version >= target_version:
            return

        # Only migrate if the field is unset or still at legacy default
        field = "hybrid_disable_threshold_mwh"
        legacy_default = 9999.0
        new_default = 20.0
        state_key = f"dc_inputs.{field}"
        current_val = st.session_state.get(state_key, dc_inputs.get(field))
        try:
            current_num = float(current_val)
        except Exception:
            current_num = None
        if current_val is None:
            is_legacy = True
        elif current_num is None or not np.isfinite(current_num):
            is_legacy = True
        else:
            is_legacy = abs(current_num - legacy_default) < 1e-6

        if is_legacy:
            st.session_state[state_key] = new_default
            dc_inputs[field] = new_default

        poi_raw = st.session_state.get("dc_inputs.poi_energy_req_mwh", dc_inputs.get("poi_energy_req_mwh"))
        if poi_raw is None:
            st.session_state[version_key] = target_version
            return

        # Align auto-default toggles with threshold intent (UI-only)
        try:
            poi_val = float(poi_raw)
        except Exception:
            poi_val = 0.0
        try:
            threshold_val = float(st.session_state.get(state_key, dc_inputs.get(field, new_default)))
        except Exception:
            threshold_val = new_default
        auto_hybrid = not (threshold_val > 0 and poi_val >= threshold_val)
        auto_cabinet_only = auto_hybrid

        def _set_toggle(name: str, value: bool) -> None:
            key = f"dc_inputs.{name}"
            st.session_state[key] = bool(value)
            dc_inputs[name] = bool(value)

        for name, auto_val in (("enable_hybrid", auto_hybrid), ("enable_cabinet_only", auto_cabinet_only)):
            current_toggle = st.session_state.get(f"dc_inputs.{name}", dc_inputs.get(name))
            if current_toggle is None:
                _set_toggle(name, auto_val)
            elif not auto_val and bool(current_toggle):
                # Flip legacy-true defaults to the new auto-off behavior
                _set_toggle(name, False)

        st.session_state[version_key] = target_version
    _migrate_dc_defaults()

    # Inject CSS
    inject_css()
    
    # Title
    st.markdown(
        """
        <div style="padding-top:0.6rem; padding-bottom:0.6rem;">
            <h1 class="calb-page-title">Utility-Scale ESS Sizing Tool V1.0 (DC)</h1>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("<br/>", unsafe_allow_html=True)
    
    # Load Data
    try:
        defaults, df_blocks, df_soh_profile, df_soh_curve, df_rte_profile, df_rte_curve = load_data(DC_DATA_PATH)
    except Exception as exc:
        st.error(f"Failed to load data file: {exc}")
        return

    if DC_DATA_IS_LEGACY:
        st.warning("using legacy dictionary")

    rte_issues = validate_rte_monotonicity_314(df_rte_profile, df_rte_curve)
    if not rte_issues.empty:
        st.warning(f"RTE curve monotonicity check found {len(rte_issues)} issues.")
        with st.expander("Show RTE monotonicity issues", expanded=False):
            st.dataframe(rte_issues, use_container_width=True)

    st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
    st.subheader("Load DC Run")
    with st.expander("Load DC Run by ID", expanded=False):
        load_run_id = st.text_input("Run ID", key="dc_load_run_id")
        if st.button("Load Run", key="dc_load_run_btn"):
            run_id = (load_run_id or "").strip()
            if not run_id:
                st.warning("Please enter a run id.")
            else:
                try:
                    with session_scope() as session:
                        access = AccessControlService(session, auth_user)
                        bundle = access.load_dc_run_bundle(run_id)
                    if not bundle:
                        st.error(f"Run ··{run_id[-8:] if len(run_id) >= 8 else run_id} not found.")
                    else:
                        case_input = (
                            bundle.input_snapshot.payload.get("case_input")
                            if bundle.input_snapshot and isinstance(bundle.input_snapshot.payload, dict)
                            else {}
                        )
                        poi_nominal_voltage_kv = case_input.get("poi_nominal_voltage_kv") or st.session_state.get(
                            "poi_nominal_voltage_kv", 33.0
                        )
                        poi_frequency_hz = case_input.get("poi_frequency_hz") or st.session_state.get(
                            "poi_frequency_hz"
                        )

                        st.session_state["poi_nominal_voltage_kv"] = poi_nominal_voltage_kv
                        st.session_state["poi_frequency_hz"] = poi_frequency_hz
                        restore_run_bundle_to_session(bundle, run_id)
                        set_run_time("dc_results")
                        st.success(f"DC run ··{run_id[-8:] if len(run_id) >= 8 else run_id} loaded.")
                except PermissionError:
                    st.error("You do not have access to this run.")
                except Exception as exc:
                    st.error(f"Failed to load run: {exc}")
    st.markdown("</div>", unsafe_allow_html=True)

    def get_default_numeric(key, fallback):
        return to_float(defaults.get(key, fallback), fallback)

    def get_default_str(key, fallback):
        return str(defaults.get(key, fallback))
        
    def get_default_percent_val(key, fallback):
         raw = to_float(defaults.get(key, fallback), fallback)
         if raw <= 1.5 and raw > 0: return raw * 100.0
         return raw

    def _init_input(field: str, default_value):
        key = f"dc_inputs.{field}"
        if key not in st.session_state:
            st.session_state[key] = default_value
        if field not in dc_inputs:
            dc_inputs[field] = st.session_state[key]
        return key

    def _init_int_input(field: str, default_value: int):
        key = _init_input(field, int(default_value))
        st.session_state[key] = to_int(st.session_state.get(key, default_value), int(default_value))
        dc_inputs[field] = st.session_state[key]
        return key

    # --- UI Form ---
    with st.container():
        st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
        st.subheader("1. Project Inputs")

        with st.form("main_form"):
            project_name_default = active_project_name or get_default_str("project_name", "CALB ESS Project")
            project_name = st.text_input(
                "Project Name",
                key=_init_input("project_name", project_name_default),
                disabled=True,
                help="Project name is controlled by the active Project selection.",
            )
            dc_inputs["project_name"] = project_name
            st.session_state["project_name"] = project_name

            c1, c2, c3 = st.columns(3)
            poi_power = c1.number_input(
                "POI Required Power (MW)",
                key=_init_input("poi_power_req_mw", get_default_numeric("poi_power_req_mw", 100.0)),
            )
            dc_inputs["poi_power_req_mw"] = poi_power
            poi_energy = c2.number_input(
                "POI Required Capacity (MWh)",
                key=_init_input("poi_energy_req_mwh", get_default_numeric("poi_energy_req_mwh", 400.0)),
            )
            dc_inputs["poi_energy_req_mwh"] = poi_energy
            project_life_key = _init_int_input(
                "project_life_years",
                int(get_default_numeric("project_life_years", 20)),
            )
            project_life = c3.number_input(
                "Project Life (Years)",
                key=project_life_key,
                step=1,
                format="%d",
            )
            project_life = to_int(project_life, 0)
            dc_inputs["project_life_years"] = project_life

            mv_default = dc_inputs.get("poi_nominal_voltage_kv")
            if mv_default is None:
                mv_default = st.session_state.get("poi_nominal_voltage_kv", 33.0)
            poi_nominal_voltage_kv = st.number_input(
                "POI / MV Voltage (kV)",
                key=_init_input("poi_nominal_voltage_kv", float(mv_default)),
                step=0.1,
            )
            dc_inputs["poi_nominal_voltage_kv"] = poi_nominal_voltage_kv

            freq_options = ["TBD", 50.0, 60.0]
            freq_default = dc_inputs.get("poi_frequency_option")
            if freq_default is None:
                freq_default = "TBD"
            if freq_default in (50, 50.0):
                freq_index = 1
            elif freq_default in (60, 60.0):
                freq_index = 2
            else:
                freq_index = 0
            poi_frequency = st.selectbox(
                "POI Frequency (Hz)",
                freq_options,
                index=freq_index,
                key=_init_input("poi_frequency_option", freq_default),
                help="Optional; used for reporting only.",
            )
            dc_inputs["poi_frequency_option"] = poi_frequency

            c4, c5, c6 = st.columns(3)
            cycles_year_key = _init_int_input(
                "cycles_per_year",
                int(get_default_numeric("cycles_per_year", 365)),
            )
            cycles_year = c4.number_input(
                "Cycles Per Year",
                key=cycles_year_key,
                step=1,
                format="%d",
            )
            cycles_year = to_int(cycles_year, 0)
            dc_inputs["cycles_per_year"] = cycles_year
            guarantee_year_key = _init_int_input("poi_guarantee_year", 0)
            guarantee_year = c5.number_input(
                "POI Guarantee Year",
                key=guarantee_year_key,
                step=1,
                format="%d",
            )
            guarantee_year = to_int(guarantee_year, 0)
            dc_inputs["poi_guarantee_year"] = guarantee_year
            sc_time_key = _init_int_input("sc_time_months", 3)
            sc_time_months = c6.number_input(
                "S&C Time (Months)",
                key=sc_time_key,
                step=1,
                format="%d",
            )
            sc_time_months = to_int(sc_time_months, 3)
            dc_inputs["sc_time_months"] = sc_time_months

            st.markdown("---")
            st.subheader("2. DC Parameters")
            
            c7, c8, c9, c10 = st.columns([1, 1, 1, 1.15])
            dod_pct = c7.number_input(
                "DOD (%)",
                key=_init_input("dod_pct", get_default_percent_val("dod_pct", 95.0)),
            )
            dc_inputs["dod_pct"] = dod_pct
            dc_rte_pct = c8.number_input(
                "DC RTE (%) System level Base on 0.25C",
                key=_init_input(
                    "dc_round_trip_efficiency_pct",
                    get_default_percent_val("dc_round_trip_efficiency_pct", 94.0),
                ),
            )
            dc_inputs["dc_round_trip_efficiency_pct"] = dc_rte_pct
            rte_adjust_pp = c9.number_input(
                "RTE Curve Adjustment (pp)",
                key=_init_input("rte_curve_adjust_pp", get_default_numeric("rte_curve_adjust_pp", 0.0)),
                min_value=-5.0,
                max_value=2.0,
                step=0.1,
            )
            dc_inputs["rte_curve_adjust_pp"] = rte_adjust_pp
            c10.markdown("<div style=\"margin-top: 1.5rem;\"></div>", unsafe_allow_html=True)
            rte_monotonic_enforce = c10.checkbox(
                "RTE Monotonic (No Rise vs SOH)",
                key=_init_input("rte_monotonic_enforce", True),
                help="Default: On. Enforces a monotonic RTE envelope so RTE does not increase as SOH declines.",
            )
            dc_inputs["rte_monotonic_enforce"] = rte_monotonic_enforce
            
            st.info(f"Design Rule: Max 418kWh Cabinets per DC Busbar (K) = {K_MAX_FIXED} (fixed)")

            st.markdown("**3. Configuration Options**")
            threshold_key = _init_input("hybrid_disable_threshold_mwh", 20.0)
            threshold_val = to_float(st.session_state.get(threshold_key, 20.0), 20.0)
            auto_enable_hybrid = not (threshold_val > 0 and poi_energy >= threshold_val)
            auto_enable_cabinet_only = auto_enable_hybrid
            copt1, copt2, copt3 = st.columns([2, 2, 3])
            enable_hybrid = copt1.checkbox(
                "Enable Hybrid Mode",
                key=_init_input("enable_hybrid", auto_enable_hybrid),
            )
            dc_inputs["enable_hybrid"] = enable_hybrid
            enable_cabinet_only = copt2.checkbox(
                "Enable Cabinet-Only Mode",
                key=_init_input("enable_cabinet_only", auto_enable_cabinet_only),
            )
            dc_inputs["enable_cabinet_only"] = enable_cabinet_only
            hybrid_disable_threshold = copt3.number_input(
                "Disable Hybrid Threshold (MWh)",
                key=threshold_key,
            )
            dc_inputs["hybrid_disable_threshold_mwh"] = hybrid_disable_threshold

            with st.expander("Advanced: Efficiency Chain"):
                poi_is_dc_side = st.checkbox(
                    "POI Is Located At DC Side (Force 100%)",
                    key=_init_input("poi_is_dc_side", False),
                )
                dc_inputs["poi_is_dc_side"] = poi_is_dc_side
                eff_dc_cables = st.number_input(
                    "DC Cables (%)",
                    key=_init_input("eff_dc_cables", 99.5),
                )
                dc_inputs["eff_dc_cables"] = eff_dc_cables
                eff_pcs = st.number_input(
                    "PCS (%)",
                    key=_init_input("eff_pcs", 98.5),
                )
                dc_inputs["eff_pcs"] = eff_pcs
                eff_mvt = st.number_input(
                    "MV Transformer (%)",
                    key=_init_input("eff_mvt", 99.5),
                )
                dc_inputs["eff_mvt"] = eff_mvt
                eff_ac_sw = st.number_input(
                    "AC Cables + SW (%)",
                    key=_init_input("eff_ac_cables_sw_rmu", 99.2),
                )
                dc_inputs["eff_ac_cables_sw_rmu"] = eff_ac_sw
                eff_hvt = st.number_input(
                    "HVT (%)",
                    key=_init_input("eff_hvt_others", 100.0),
                )
                dc_inputs["eff_hvt_others"] = eff_hvt

            st.markdown("---")
            run_btn = st.form_submit_button("Run Sizing")
        
        st.markdown("</div>", unsafe_allow_html=True)

    # --- Logic Execution ---
    if run_btn:
        bump_run_id_dc()
        poi_mv_kv = float(poi_nominal_voltage_kv)
        st.session_state["poi_nominal_voltage_kv"] = poi_mv_kv
        st.session_state["grid_kv"] = poi_mv_kv
        ac_inputs["grid_kv"] = poi_mv_kv
        ac_inputs["mv_kv"] = poi_mv_kv
        poi_frequency_hz = None if poi_frequency == "TBD" else float(poi_frequency)
        st.session_state["poi_frequency_hz"] = poi_frequency_hz
        dc_inputs["poi_frequency_hz"] = poi_frequency_hz
        dc_inputs["poi_nominal_voltage_kv"] = poi_mv_kv
        dc_inputs["project_name"] = project_name
        dc_inputs["poi_power_req_mw"] = poi_power
        dc_inputs["poi_energy_req_mwh"] = poi_energy
        
        if poi_is_dc_side:
            eff_dc_cables = eff_pcs = eff_mvt = eff_ac_sw = eff_hvt = 100.0

        case_frequency_hz = 50.0 if poi_frequency_hz is None else float(poi_frequency_hz)
        case_input_base = {
            "project_name": project_name,
            "poi_power_req_mw": poi_power,
            "poi_energy_req_mwh": poi_energy,
            "poi_nominal_voltage_kv": poi_mv_kv,
            "poi_frequency_hz": case_frequency_hz,
            "project_life_years": project_life,
            "cycles_per_year": cycles_year,
            "poi_guarantee_year": guarantee_year,
            "eff_dc_cables": eff_dc_cables,
            "eff_pcs": eff_pcs,
            "eff_mvt": eff_mvt,
            "eff_ac_cables_sw_rmu": eff_ac_sw,
            "eff_hvt_others": eff_hvt,
            "sc_time_months": sc_time_months,
            "dod_pct": dod_pct,
            "dc_round_trip_efficiency_pct": dc_rte_pct,
            "rte_curve_adjust_pp": rte_adjust_pp,
            "rte_monotonic_enforce": rte_monotonic_enforce,
        }

        sc_loss_pct = calc_sc_loss_pct(sc_time_months)
        inputs = {
            "project_name": project_name,
            "poi_power_req_mw": poi_power,
            "poi_energy_req_mwh": poi_energy,
            "eff_dc_cables": eff_dc_cables,
            "eff_pcs": eff_pcs,
            "eff_mvt": eff_mvt,
            "eff_ac_cables_sw_rmu": eff_ac_sw,
            "eff_hvt_others": eff_hvt,
            "sc_time_months": sc_time_months,
            "sc_loss_frac": sc_loss_pct / 100.0,
            "dod_pct": dod_pct,
            "dc_round_trip_efficiency_pct": dc_rte_pct,
            "rte_curve_adjust_pp": rte_adjust_pp,
            "rte_monotonic_enforce": rte_monotonic_enforce,
            "project_life_years": project_life,
            "cycles_per_year": cycles_year,
            "poi_guarantee_year": guarantee_year
        }

        s1 = run_stage1(inputs, defaults)

        # Stage 1 Display
        st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
        st.subheader("3. Stage 1 - DC Requirement")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Theoretical DC Required", f"{s1['dc_energy_capacity_required_mwh']:.2f} MWh")
        m2.metric("Efficiency Chain", f"{s1['eff_dc_to_poi_frac']*100:.2f} %")
        m3.metric("DC Usable @BOL", f"{s1['dc_usable_bol_frac']*100:.2f} %")
        m4.metric("DC Power Req", f"{s1['dc_power_required_mw']:.2f} MW")
        rte_base_pct = float(s1.get("dc_rte_base_frac", s1.get("dc_round_trip_efficiency_frac", 0.0))) * 100.0
        rte_adjust_pp_used = float(s1.get("rte_curve_adjust_pp", 0.0))
        rte_effective_pct = float(s1.get("dc_rte_effective_frac", s1.get("dc_round_trip_efficiency_frac", 0.0))) * 100.0
        r1, r2, r3 = st.columns(3)
        r1.metric("DC RTE base (%)", f"{rte_base_pct:.2f} %")
        r2.metric("RTE adjustment (pp)", f"{rte_adjust_pp_used:+.2f} pp")
        r3.metric("DC RTE effective (%)", f"{rte_effective_pct:.2f} %")
        st.markdown("</div>", unsafe_allow_html=True)

        # Run Options
        modes_to_run = ["container_only"]
        if enable_cabinet_only:
            modes_to_run.insert(0, "cabinet_only")
        if enable_hybrid:
            modes_to_run.insert(0, "hybrid")

        results = {}
        for mode in modes_to_run:
            try:
                results[mode] = size_with_guarantee(
                    s1, mode,
                    df_blocks,
                    df_soh_profile, df_soh_curve,
                    df_rte_profile, df_rte_curve,
                    k_max=K_MAX_FIXED
                )
            except Exception as e:
                results[mode] = ("ERROR", str(e))

        ok_results = {
            k: v for k, v in results.items() if isinstance(v, tuple) and v[0] != "ERROR"
        }
        report_order = [
            (k, k.replace("_", " ").title()) for k in modes_to_run if k in ok_results
        ]
        dc_results["results_dict"] = ok_results
        dc_results["report_order"] = report_order
        active_mode = first_success_key(ok_results, modes_to_run)
        active_snapshot = None
        if active_mode:
            try:
                bundle = DcExcelMasterDataBundle(
                    workbook_path=Path(DC_DATA_PATH),
                    defaults=defaults,
                    df_blocks=df_blocks.copy(),
                    df_soh_profile=df_soh_profile.copy(),
                    df_soh_curve=df_soh_curve.copy(),
                    df_rte_profile=df_rte_profile.copy(),
                    df_rte_curve=df_rte_curve.copy(),
                    raw_sheets={},
                )
                active_snapshot = service_size_with_guarantee(
                    Stage1Result.model_validate(s1),
                    active_mode,
                    bundle,
                    k_max=K_MAX_FIXED,
                )
                case_input = dict(case_input_base)
                case_input["scenario_id"] = active_mode
                persist_result = persist_dc_run(
                    SizingCaseInput.model_validate(case_input),
                    active_snapshot,
                    project_id=active_project_id,
                    sizing_case_id=active_case_id,
                    case_code=active_case_code,
                    case_name=active_case_name,
                    defaults=defaults,
                    source_ref="dc_view",
                    actor=auth_user.username,
                )
                run_id = persist_result.get("run_id")
                if run_id:
                    with session_scope() as session:
                        access = AccessControlService(session, auth_user)
                        bundle = access.load_dc_run_bundle(run_id)
                    if not bundle:
                        st.error("Run persisted but could not be reloaded from DB.")
                    else:
                        st.session_state["poi_nominal_voltage_kv"] = poi_nominal_voltage_kv
                        st.session_state["poi_frequency_hz"] = poi_frequency_hz
                        restore_run_bundle_to_session(bundle, run_id)
                        st.info("DC run saved and restored successfully.")
            except Exception as exc:
                st.error(f"Failed to persist DC run: {exc}")

        # Tabs Display
        st.markdown("<div class='calb-card'>", unsafe_allow_html=True)
        st.subheader("4. Configuration Comparison")
        
        tabs = st.tabs([m.replace("_", " ").title() for m in modes_to_run])
        for i, mode in enumerate(modes_to_run):
            with tabs[i]:
                res = results.get(mode)
                if isinstance(res, tuple) and res[0] != "ERROR":
                    s2, s3_df, s3_meta, iter_c, poi_g, conv = res
                    
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Installed DC Nameplate", f"{s2['dc_nameplate_bol_mwh']:.3f} MWh")
                    c2.metric("Container Count", int(s2.get('container_count', 0)))
                    c3.metric("Cabinet Count", int(s2.get('cabinet_count', 0)))

                    st.dataframe(s2.get("block_config_table"), use_container_width=True)
                    
                    # Chart
                    s3_df_sorted = s3_df.sort_values("Year_Index").reset_index(drop=True)
                    bars = alt.Chart(s3_df_sorted).mark_bar(color=CALB_SKY_BLUE).encode(
                        x=alt.X("Year_Index:O", title="Year"),
                        y=alt.Y("POI_Usable_Energy_MWh:Q", title="POI Usable (MWh)"),
                        tooltip=["Year_Index", "POI_Usable_Energy_MWh", "SOH_Display_Pct"]
                    )
                    rule = alt.Chart(pd.DataFrame({"y": [float(poi_energy)]})).mark_rule(color="red").encode(y="y")
                    st.altair_chart((bars + rule).properties(height=350), use_container_width=True)

                    # --- RESTORED YEARLY DATA TABLE ---
                    with st.expander("Show yearly data table", expanded=False):
                        st.markdown(
                            """
<style>
/* Yearly table: left-align headers/cells and hide index column */
div[data-testid="stDataFrame"] div[role="columnheader"],
div[data-testid="stDataFrame"] div[role="gridcell"],
div[data-testid="stDataFrame"] div[role="columnheader"] div,
div[data-testid="stDataFrame"] div[role="gridcell"] div {
  justify-content: flex-start !important;
  text-align: left !important;
}
div[data-testid="stDataFrame"] div[role="rowheader"] {
  display: none !important;
}
</style>
""",
                            unsafe_allow_html=True,
                        )

                        def _col_or_nan(df: pd.DataFrame, col: str):
                            if col in df.columns:
                                return df[col]
                            return pd.Series([np.nan] * len(df), index=df.index)

                        year_series = _col_or_nan(s3_df_sorted, "Year_Index")
                        year_numeric = pd.to_numeric(year_series, errors="coerce")
                        if year_numeric.isna().all():
                            year_numeric = pd.Series(range(len(s3_df_sorted)), index=s3_df_sorted.index)

                        def _fmt_bool_text(val):
                            if pd.isna(val):
                                return ""
                            return "Yes" if bool(val) else "No"

                        disp_show = pd.DataFrame({
                            "Year": year_numeric.astype("Int64"),
                            "DC Usable (MWh)": pd.to_numeric(_col_or_nan(s3_df_sorted, "DC_Usable_MWh"), errors="coerce").round(3),
                            "POI Usable (MWh)": pd.to_numeric(_col_or_nan(s3_df_sorted, "POI_Usable_Energy_MWh"), errors="coerce").round(3),
                            "DC RTE (%)": pd.to_numeric(_col_or_nan(s3_df_sorted, "DC_RTE_Pct"), errors="coerce").round(3),
                            "System RTE (%)": pd.to_numeric(_col_or_nan(s3_df_sorted, "System_RTE_Pct"), errors="coerce").round(3),
                            "Meets Req?": _col_or_nan(s3_df_sorted, "Meets_POI_Req").apply(_fmt_bool_text),
                            "Guarantee Year": _col_or_nan(s3_df_sorted, "Is_Guarantee_Year").apply(_fmt_bool_text),
                        }).reset_index(drop=True)

                        # Display table with CSS-driven left alignment (sorting stays numeric where possible)
                        try:
                            st.dataframe(disp_show, use_container_width=True, hide_index=True)
                        except TypeError:
                            # Fallback for older Streamlit versions without hide_index
                            disp_show.index = [""] * len(disp_show)
                            st.dataframe(disp_show, use_container_width=True)
                    # ----------------------------------

                    # Pack data for Session State (For AC/SLD)
                    # We only pack the FIRST valid result as the 'active' one for now, or the user preferred one
                    if active_mode and mode == active_mode:
                        container_count = int(s2.get("container_count", 0))
                        cabinet_count = int(s2.get("cabinet_count", 0))
                        total_blocks = container_count + cabinet_count
                        total_mwh = float(s2.get("dc_nameplate_bol_mwh", 0.0))

                        # Prefer explicit unit capacity for container-only; otherwise use weighted average.
                        block_unit_mwh = 0.0
                        if cabinet_count == 0 and container_count > 0:
                            block_unit_mwh = total_mwh / container_count if container_count else 0.0
                            block_table = s2.get("block_config_table")
                            if block_table is not None and not block_table.empty:
                                try:
                                    block_unit_mwh = float(block_table.iloc[0]["Unit Capacity (MWh)"])
                                except Exception:
                                    pass
                        if block_unit_mwh <= 0:
                            block_unit_mwh = (
                                total_mwh / total_blocks if total_blocks > 0 else get_standard_container_mwh()
                            )

                        dc_res = DCBlockResult(
                            block_id="DC-Block",
                            container_model="CALB-314Ah",  # Ideal: read from config
                            capacity_mwh=block_unit_mwh,
                            voltage_v=1200.0,
                            count=total_blocks,
                        )

                        st.session_state["dc_result_summary"] = {
                            "mwh": total_mwh,
                            "target_mw": poi_power,
                            "voltage": 1200,
                            "container_count": container_count,
                            "cabinet_count": cabinet_count,
                            "total_blocks": total_blocks,
                            "mode": active_mode,
                            "dc_block": dc_res,  # Pass the object for AC view
                        }

                        # Also pack for legacy interface if needed
                        st.session_state["stage13_output"] = pack_stage13_output(
                            s1,
                            s2,
                            s3_meta,
                            dc_block_total_qty=total_blocks,
                            selected_scenario=active_mode,
                            poi_nominal_voltage_kv=poi_nominal_voltage_kv,
                            poi_frequency_hz=poi_frequency_hz,
                            stage3_df=s3_df,
                        )
                        dc_results.update(
                            {
                                "stage13_output": st.session_state.get("stage13_output"),
                                "dc_result_summary": st.session_state.get("dc_result_summary"),
                            }
                        )
                        set_run_time("dc_results")

                else:
                    st.error(f"Error: {res[1]}")
        st.markdown("</div>", unsafe_allow_html=True)

        # Export Button
        if DOCX_AVAILABLE:
            ok_results = {k: v for k, v in results.items() if v[0] != "ERROR"}
            report_order = [(k, k.replace("_", " ").title()) for k in modes_to_run if k in ok_results]
            report_bytes = build_report_bytes(s1, ok_results, report_order)
            
            if report_bytes:
                st.download_button(
                    "Export Technical Sizing Report",
                    data=report_bytes,
                    file_name=make_report_filename(project_name),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        st.info("Sizing Complete. Please proceed to AC Sizing or SLD generation via sidebar.")










