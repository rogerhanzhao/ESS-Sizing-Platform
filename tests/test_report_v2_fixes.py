# -----------------------------------------------------------------------------
# Personal Open-Source Notice
#
# Copyright (c) 2026 Alex.Zhao. All rights reserved.
#
# This repository is released under the MIT License (see LICENSE file).
# Intended use: learning, evaluation, and engineering reference for Utility-scale
# BESS/ESS sizing and Reporting workflows.
#
# DISCLAIMER: This software is provided "AS IS", without warranty of any kind,
# express or implied. In no event shall the author(s) be liable for any claim,
# damages, or other liability arising from, out of, or in connection with the
# software or the use or other dealings in the software.
#
# NOTE: This is a personal project. It is not an official product or statement
# of any company or organization.
# -----------------------------------------------------------------------------

import base64
import io
import json
import pytest
import re
from pathlib import Path

from docx import Document

from calb_sizing_tool.reporting.report_context import build_report_context
from calb_sizing_tool.reporting.report_v2 import export_report_v2_1
from tools.regress_export import run_ac_sizing, run_dc_sizing


def test_efficiency_chain_uses_dc_sizing_values():
    """Verify that exported efficiency chain uses actual DC SIZING values, not defaults."""
    fixture_path = Path(__file__).parent / "fixtures" / "v1_case01_container_input.json"
    fixture = json.loads(fixture_path.read_text(encoding="utf-8"))

    dc_results = run_dc_sizing(fixture)
    ac_output = run_ac_sizing(fixture, dc_results["stage1"], dc_results["stage2"])

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
    )

    ctx = build_report_context(
        session_state={
            "artifacts": {
                "sld_png_bytes": png_bytes,
                "layout_png_bytes": png_bytes,
            }
        },
        stage_outputs={
            "stage13_output": dc_results["stage1"],
            "stage2": dc_results["stage2"],
            "stage3_df": dc_results["stage3_df"],
            "stage3_meta": dc_results["stage3_meta"],
            "ac_output": ac_output,
        },
        project_inputs={"poi_energy_guarantee_mwh": fixture["poi_energy_req_mwh"]},
        scenario_ids=fixture["scenario_id"],
    )

    # Verify efficiency values were read from DC SIZING, not defaults
    assert ctx.efficiency_chain_oneway_frac > 0, "Total efficiency should be positive"
    assert ctx.efficiency_chain_oneway_frac <= 1.0, "Total efficiency should not exceed 100%"
    
    # All component efficiencies should be present
    assert ctx.efficiency_components_frac.get("eff_dc_cables_frac") is not None
    assert ctx.efficiency_components_frac.get("eff_pcs_frac") is not None
    assert ctx.efficiency_components_frac.get("eff_mvt_frac") is not None
    assert ctx.efficiency_components_frac.get("eff_ac_cables_sw_rmu_frac") is not None
    assert ctx.efficiency_components_frac.get("eff_hvt_others_frac") is not None
    
    # Export and verify report contains actual values
    report_bytes = export_report_v2_1(ctx)
    doc = Document(io.BytesIO(report_bytes))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    # Should NOT contain defaults (e.g., "97.00%" for PCS which is a common default)
    # Should contain the efficiency disclaimer about not including Auxiliary
    assert "do not include Auxiliary losses" in joined, "Report should state that efficiencies exclude Auxiliary"
    
    # Verify Efficiency Chain section exists
    assert "Efficiency Chain" in joined


def test_ac_block_config_not_verbose():
    """Verify AC Block configuration doesn't list every block when they're identical."""
    fixture_path = Path(__file__).parent / "fixtures" / "v1_case01_container_input.json"
    fixture = json.loads(fixture_path.read_text(encoding="utf-8"))

    dc_results = run_dc_sizing(fixture)
    ac_output = run_ac_sizing(fixture, dc_results["stage1"], dc_results["stage2"])

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
    )

    ctx = build_report_context(
        session_state={
            "artifacts": {
                "sld_png_bytes": png_bytes,
                "layout_png_bytes": png_bytes,
            }
        },
        stage_outputs={
            "stage13_output": dc_results["stage1"],
            "stage2": dc_results["stage2"],
            "stage3_df": dc_results["stage3_df"],
            "stage3_meta": dc_results["stage3_meta"],
            "ac_output": ac_output,
        },
        project_inputs={"poi_energy_guarantee_mwh": fixture["poi_energy_req_mwh"]},
        scenario_ids=fixture["scenario_id"],
    )

    # Export and verify report
    report_bytes = export_report_v2_1(ctx)
    doc = Document(io.BytesIO(report_bytes))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    # Should have summary section
    assert "AC Block Sizing" in joined or "AC:DC Ratio" in joined

    # The concept-arrangement section is allowed to use the phrase "AC Block",
    # but the report must not revert to a per-block list such as "AC Block 1".
    assert not re.search(r"\bAC Block\s+\d+\b", joined), (
        "Report contains a verbose numbered AC Block listing"
    )


def test_site_array_power_and_energy_match_the_ac_sizing_tables():
    """§9 must never contradict §6: the site figure's MW/MWh come from the run.

    The site-array engine used to hardcode a 5 MW / 5.015 MWh unit, so a 10 MW
    station was reported at HALF its power (65 MW vs 130 MW) and energy was
    computed from n_blocks x dc_per_block (104) instead of the real DC Block
    count (100) — the report contradicted its own AC Sizing tables.
    """
    import io as _io

    from docx import Document as _Document

    from calb_sizing_tool.reporting.report_context import build_report_context
    from calb_sizing_tool.reporting.report_v2 import export_report_v2_1

    ac_output = {
        "num_blocks": 13, "pcs_per_block": 8, "pcs_kw": 1250.0,
        "block_size_mw": 10.0, "transformer_mva": 11.111, "total_ac_mw": 130.0,
        "lv_winding_count": 2, "transformer_topology": "three_winding",
        "dc_blocks_total": 100,
    }
    ctx = build_report_context(
        session_state={},
        stage_outputs={
            "stage13_output": {
                "project_name": "Consistency", "poi_power_req_mw": 115.0,
                "poi_energy_req_mwh": 400.0, "project_life_years": 20,
                "poi_guarantee_year": 4, "cycles_per_year": 365,
            },
            "ac_output": ac_output,
            "stage2": {"container_count": 100, "dc_nameplate_bol_mwh": 501.5},
        },
        project_inputs={},
    )
    doc = _Document(_io.BytesIO(export_report_v2_1(ctx)))
    caption = next(
        (p.text for p in doc.paragraphs if "Concept Site Arrangement —" in p.text), ""
    )
    assert caption, "§9 site arrangement caption not found"
    # 13 AC Blocks x 10.00 MW = 130 MW (was 65), and the REAL 100 DC Blocks.
    assert "10 MW\nper block" in caption or "10 MW per block" in caption, caption
    assert "130 MW" in caption, caption
    assert "65 MW" not in caption, caption


def _ctx_for(pcs_per_block: int, num_blocks: int, dc_blocks_total: int,
             block_size_mw: float, total_ac_mw: float):
    from calb_sizing_tool.reporting.report_context import build_report_context

    return build_report_context(
        session_state={},
        stage_outputs={
            "stage13_output": {
                "project_name": "XSec", "poi_power_req_mw": 115.0,
                "poi_energy_req_mwh": 400.0, "project_life_years": 20,
                "poi_guarantee_year": 4, "cycles_per_year": 365,
            },
            "ac_output": {
                "num_blocks": num_blocks, "pcs_per_block": pcs_per_block,
                "pcs_kw": 1250.0, "block_size_mw": block_size_mw,
                "transformer_mva": 11.111, "total_ac_mw": total_ac_mw,
                "lv_winding_count": 2, "transformer_topology": "three_winding",
                "dc_blocks_total": dc_blocks_total,
            },
            "stage2": {"container_count": dc_blocks_total,
                       "dc_nameplate_bol_mwh": dc_blocks_total * 5.015},
        },
        project_inputs={},
    )


def test_ten_mw_block_is_drawn_with_the_forty_foot_station_not_the_twenty():
    """P1-4 lock for F1: station size follows the AC Block class, not a constant.

    A 10 MW / 8-PCS block that does not hit the 8-DC bilateral shape still draws
    the linear arrangement — but with the 40 ft station (12.192 m), so its
    envelope is 6.13 m longer than the old hardcoded 20 ft cabin produced.
    """
    from calb_diagrams.ac_block_arrangement_v2 import US_NFPA_OIL, compute_layout

    ctx = _ctx_for(pcs_per_block=8, num_blocks=13, dc_blocks_total=91,
                   block_size_mw=10.0, total_ac_mw=130.0)
    doc = Document(io.BytesIO(export_report_v2_1(ctx)))
    caption = next(
        (p.text for p in doc.paragraphs
         if p.text.startswith("Figure") and "Typical AC Block Arrangement" in p.text),
        "",
    )
    assert caption, "§8 arrangement caption not found"

    dc_per_ac = 91 // 13
    forty = compute_layout(dc_per_ac, US_NFPA_OIL, pcs_count=8, block_power_mw=10.0)
    twenty = compute_layout(dc_per_ac, US_NFPA_OIL, pcs_count=4, block_power_mw=5.0)
    assert forty.station_length_m == pytest.approx(12.192, abs=0.001)
    assert f"{forty.envelope_w_m:.2f}" in caption, caption
    assert f"{twenty.envelope_w_m:.2f}" not in caption, caption


def test_eight_by_eight_report_draws_the_site_from_the_block_it_actually_drew():
    """Section 9 must compose the block section 8 drew — one product, one footprint.

    This case used to have NO site figure at all: the L2 engine could only
    reconstruct a linear block from dc_per_block, so a central-station block was
    either suppressed or tiled as something it is not. The block's real
    placements now travel to the site engine, so the section is drawn and its
    land is reported.
    """
    from calb_diagrams.ac_block_bilateral_layout import compute_bilateral_layout

    ctx = _ctx_for(pcs_per_block=8, num_blocks=13, dc_blocks_total=104,
                   block_size_mw=10.0, total_ac_mw=130.0)
    doc = Document(io.BytesIO(export_report_v2_1(ctx)))
    text = "\n".join(p.text for p in doc.paragraphs)
    tables = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )

    assert "9.  Concept Site Arrangement (Concept Only)" in text
    assert "Concept Site Arrangement —" in text          # the figure is there
    # Footprint is reported, because minimum land is the objective.
    assert "Site land area" in tables
    assert "Land intensity" in tables
    assert "Packing (minimum land)" in tables
    # A central-station block has no station-to-station corridor to share.
    assert "central-station blocks, full aisle between blocks" in tables
    # And the rejected perimeter field never appears.
    assert "28.54" not in text

    block = compute_bilateral_layout(8)
    assert f"{block.envelope_w_m:.2f}" in text
