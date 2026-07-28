"""End-to-end: a mixed AC Block station reaches the report as a head/tail schedule.

Uniform stations must NOT show the mixed schedule (one code path, but the extra
table appears only when it adds information).
"""
from __future__ import annotations

import io

from docx import Document

from calb_sizing_tool.reporting.report_context import build_report_context
from calb_sizing_tool.reporting.report_v2 import (
    _mixed_head_entry,
    _representative_dc_per_ac,
    export_report_v2_1,
)
from calb_sizing_tool.services.ac_mixed_station import summarize_ac_block_rows


def _build_ctx(ac_output: dict):
    return build_report_context(
        session_state={},
        stage_outputs={
            "stage13_output": {
                "project_name": "Mixed Station",
                "poi_power_req_mw": 45.0,
                "poi_energy_req_mwh": 180.0,
                "project_life_years": 20,
                "poi_guarantee_year": 0,
                "cycles_per_year": 365,
            },
            "ac_output": ac_output,
            "stage2": {"container_count": 36},
        },
        project_inputs={"grid_power_factor": 0.9},
    )


def _report_text(ac_output: dict) -> str:
    ctx = _build_ctx(ac_output)
    doc = Document(io.BytesIO(export_report_v2_1(ctx)))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _mixed_ac_output() -> dict:
    # 4 head blocks (10 MW, 8x1250, 8 DC) + 1 tail (5 MW, 4x1250, 4 DC) = 36 DC.
    rows = [{"pcs_count": 8, "pcs_kw": 1250, "dc_blocks": 8} for _ in range(4)]
    rows.append({"pcs_count": 4, "pcs_kw": 1250, "dc_blocks": 4})
    breakdown = summarize_ac_block_rows(rows)
    return {
        "selected_ratio": "1:8",
        "num_blocks": len(rows),
        "pcs_per_block": 8,
        "pcs_count_by_block": [r["pcs_count"] for r in rows],
        "pcs_kw": 1250,
        "block_size_mw": 10.0,
        "total_ac_mw": 45.0,
        "dc_blocks_per_ac": [r["dc_blocks"] for r in rows],
        "dc_blocks_total": 36,
        "grid_power_factor": 0.9,
        "transformer_mva": 10.0 / 0.9,
        "ac_block_mixed": True,
        "ac_block_rows": rows,
        "ac_block_breakdown": breakdown,
    }


def test_mixed_station_renders_head_and_tail_schedule():
    text = _report_text(_mixed_ac_output())
    assert "Mixed AC Block Station Schedule" in text
    assert "Head" in text and "Tail" in text
    # Tail model (5 MW, 4 x 1250) and head model (10 MW, 8 x 1250) both present.
    assert "10.00 MW" in text
    assert "5.00 MW" in text
    # The tail is tied back to the DC container + cabinet tail for coherence.
    assert "counterpart of the DC container + cabinet tail" in text


def test_uniform_station_has_no_mixed_schedule():
    rows = [{"pcs_count": 8, "pcs_kw": 1250, "dc_blocks": 8} for _ in range(5)]
    breakdown = summarize_ac_block_rows(rows)
    ac_output = {
        "selected_ratio": "1:8",
        "num_blocks": 5,
        "pcs_per_block": 8,
        "pcs_count_by_block": [8] * 5,
        "pcs_kw": 1250,
        "block_size_mw": 10.0,
        "total_ac_mw": 50.0,
        "dc_blocks_per_ac": [8] * 5,
        "dc_blocks_total": 40,
        "grid_power_factor": 0.9,
        "transformer_mva": 10.0 / 0.9,
        "ac_block_mixed": False,
        "ac_block_rows": rows,
        "ac_block_breakdown": breakdown,
    }
    text = _report_text(ac_output)
    assert "Mixed AC Block Station Schedule" not in text


def test_representative_dc_per_ac_uses_head_not_average_for_mixed():
    # 4 × 8-DC head + 1 × 4-DC tail = 36 DC over 5 blocks. The average rounds to
    # 7 (a block that exists nowhere); the representative must be the head's 8.
    ctx = _build_ctx(_mixed_ac_output())
    head = _mixed_head_entry(ctx)
    assert head is not None and head["role"] == "Head"
    assert _representative_dc_per_ac(ctx) == 8


def test_representative_dc_per_ac_is_average_for_uniform():
    rows = [{"pcs_count": 8, "pcs_kw": 1250, "dc_blocks": 8} for _ in range(4)]
    rows.append({"pcs_count": 8, "pcs_kw": 1250, "dc_blocks": 4})  # uneven DC, one model
    ac_output = {
        "selected_ratio": "1:8",
        "num_blocks": 5,
        "pcs_per_block": 8,
        "pcs_count_by_block": [8] * 5,
        "pcs_kw": 1250,
        "block_size_mw": 10.0,
        "total_ac_mw": 50.0,
        "dc_blocks_per_ac": [8, 8, 8, 8, 4],
        "dc_blocks_total": 36,
        "grid_power_factor": 0.9,
        "transformer_mva": 10.0 / 0.9,
        "ac_block_mixed": False,
        "ac_block_rows": rows,
        "ac_block_breakdown": summarize_ac_block_rows(rows),
    }
    ctx = _build_ctx(ac_output)
    assert _mixed_head_entry(ctx) is None
    assert _representative_dc_per_ac(ctx) == round(36 / 5)  # 7 (average), single model


def test_mixed_report_draws_head_block_and_states_real_whole_site_power():
    text = _report_text(_mixed_ac_output())
    # §8/§9 disclose the representative head block for a mixed station.
    assert "representative Head AC Block" in text
    # Whole-site power is the real 45 MW head + tail total, not the uniform
    # over-count of 5 blocks × 10 MW head = 50 MW.
    assert "50 MW / " not in text
