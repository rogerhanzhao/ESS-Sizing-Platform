"""Concept whole-site LAYOUT for a governed site (every block at real footprint).

Deliverable: a drawable whole-site arrangement bound to real products (not just a
band/count overview), plus the report's provisional equipment schedule sourced
from the actual AC Sizing run.
"""
from __future__ import annotations

from calb_diagrams.governed_site_layout_concept import (
    render_governed_site_layout_concept_svg,
)
from calb_sizing_tool.services.governed_ac_block_service import build_governed_site_run


def test_layout_places_every_governed_block_and_marks_concept():
    run = build_governed_site_run(92)  # 11x8 + 1x4
    svg = render_governed_site_layout_concept_svg(run)
    assert svg.startswith("<svg") and svg.endswith("</svg>")
    # A block rect per placed AC block (12 governed AC blocks) plus ground rects.
    assert svg.count("<rect") >= run.ac_blocks_total
    assert "CONCEPT ONLY" in svg and "NOT FOR CONSTRUCTION" in svg
    # Site envelope estimate is stated.
    assert "site envelope" in svg


def test_pure_multiple_of_eight_single_family_layout():
    run = build_governed_site_run(16)  # 2 x 8
    svg = render_governed_site_layout_concept_svg(run)
    assert svg.startswith("<svg")
    assert "20.00 MW" in svg


def test_typed_api_and_explicit_params():
    from calb_diagrams.governed_site_layout_concept import (
        ConceptSiteLayoutParameters,
        concept_site_layout_from_run,
        render_concept_site_layout_svg,
    )

    run = build_governed_site_run(92)
    layout = concept_site_layout_from_run(run)
    # Strongly typed: blocks carry their own fields (no ac_output duck-typing).
    assert layout.ac_blocks_total == 12
    codes = [b.configuration_code for b in layout.blocks]
    assert codes == [
        "ACBLK-10MW-8PCS-8DC-40FT-BILATERAL",
        "ACBLK-5MW-4PCS-4DC-40FT-LINEAR",
    ]
    # Packing constants are explicit parameters, not module globals.
    params = ConceptSiteLayoutParameters(site_width_m=200.0, aisle_m=8.0, row_gap_m=12.0)
    svg = render_concept_site_layout_svg(layout, params=params, title="SITE")
    assert svg.startswith("<svg") and svg.endswith("</svg>")
    assert "CONCEPT ONLY" in svg
    assert "8 m aisle" in svg  # reflects the explicit param


def test_report_equipment_schedule_and_layout_for_governed_run():
    import io

    from docx import Document

    from calb_sizing_tool.reporting.report_context import build_report_context
    from calb_sizing_tool.reporting.report_v2 import export_report_v2_1
    from calb_sizing_tool.services.governed_ac_block_service import (
        build_governed_primary_ac_output,
    )

    ac = build_governed_primary_ac_output(92)
    ctx = build_report_context(
        session_state={},
        stage_outputs={
            "stage13_output": {
                "project_name": "Mix",
                "poi_power_req_mw": 115.0,
                "poi_energy_req_mwh": 400.0,
                "project_life_years": 20,
                "poi_guarantee_year": 0,
                "cycles_per_year": 365,
            },
            "ac_output": ac,
            "stage2": {"container_count": 92},
        },
        project_inputs={},
    )
    doc = Document(io.BytesIO(export_report_v2_1(ctx)))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(parts)
    # Equipment schedule section with both governed families and a TBD transformer
    # (settings-only run) — never a MW/PF fabrication.
    assert "Equipment Schedule" in text
    assert "ACBLK-10MW-8PCS-8DC-40FT-BILATERAL" in text
    assert "ACBLK-5MW-4PCS-4DC-40FT-LINEAR" in text
    assert "TBD" in text
    assert "÷ 0.9" not in text and "/ 0.9" not in text
