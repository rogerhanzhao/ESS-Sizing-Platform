"""Bind a real catalogue product to the auto-recommend (trunk) AC Block spec.

The generic flow computes a PCS count × rating spec; when a catalogue product
matches that spec the flow can bind it for a confirmed transformer nameplate /
vector group instead of the MW ÷ PF estimate. No new parameter — a lookup by the
spec already computed.
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from calb_sizing_tool.infra.db.base import Base
from calb_sizing_tool.infra.db.session import create_engine_for_url
from calb_sizing_tool.reporting.report_context import build_report_context
from calb_sizing_tool.reporting.report_v2 import export_report_v2_1
from calb_sizing_tool.services.ac_block_product_seed_service import seed_ac_block_products
from calb_sizing_tool.services.ac_block_product_match import (
    match_ac_block_products,
    preferred_catalogue_architecture,
    product_transformer_overrides,
)


@pytest.fixture()
def seeded_db(tmp_path) -> str:
    db_url = f"sqlite:///{(tmp_path / 'products.sqlite').as_posix()}"
    Base.metadata.create_all(bind=create_engine_for_url(db_url))
    seed_ac_block_products(db_url=db_url)
    return db_url


def test_matches_same_pcs_combo_only(seeded_db):
    # 4 x 1250 kW (5 MW) -> Kehua BCS5000K variants (same combo).
    codes = {m["block_code"] for m in match_ac_block_products(4, 1250.0, db_url=seeded_db)}
    assert codes and all("5000K" in c for c in codes)
    # 2 x 2500 kW is a different combo (NR-5000), not returned for 4 x 1250.
    assert not any("9567MV-5000" in c for c in codes)
    # 2 x 1250 -> 2.5 MW Kehua.
    assert all("2500K" in m["block_code"] for m in match_ac_block_products(2, 1250.0, db_url=seeded_db))
    # A spec with no catalogue product returns nothing (no fabrication).
    assert match_ac_block_products(3, 1725.0, db_url=seeded_db) == []


def test_invalid_spec_returns_empty(seeded_db):
    assert match_ac_block_products(0, 1250.0, db_url=seeded_db) == []
    assert match_ac_block_products(4, 0, db_url=seeded_db) == []


def test_product_preference_is_contextual_and_does_not_change_matching_spec(seeded_db):
    default = preferred_catalogue_architecture("1:4")
    one_to_eight = preferred_catalogue_architecture("1:8")
    assert (default["pcs_count"], default["pcs_kw"]) == (2, 2500)
    assert (one_to_eight["pcs_count"], one_to_eight["pcs_kw"]) == (8, 1250)

    # The 1:8 small-PCS candidates are returned only after the user selected
    # 8 x 1250 and the compatible three-winding topology; no call substitutes
    # that architecture for a different selected PCS combination.
    matches = match_ac_block_products(
        8,
        1250.0,
        grouping_ratio="1:8",
        transformer_topology="three_winding",
        db_url=seeded_db,
    )
    assert {m["block_code"] for m in matches} >= {
        "SINENG-EH-10000-HB-UD-10-33",
        "KEHUA-BCS10000K-C-HUD-T8",
    }
    assert all(m["transformer_topology"] in (None, "three_winding") for m in matches)
    assert any(m["is_preferred_architecture"] for m in matches)
    assert match_ac_block_products(
        8,
        1250.0,
        grouping_ratio="1:8",
        transformer_topology="two_winding",
        db_url=seeded_db,
    ) == []


def test_overrides_carry_only_real_values(seeded_db):
    m = match_ac_block_products(4, 1250.0, db_url=seeded_db)[0]
    ov = product_transformer_overrides(m)
    assert ov["transformer_mva"] == 5.0
    assert ov["ac_block_product_block_code"] == m["block_code"]
    # An empty product yields no override (never a fabricated value).
    assert product_transformer_overrides(None) == {}


def test_report_shows_real_nameplate_for_a_bound_generic_run():
    # A GENERIC (non-governed) run that bound a catalogue product must show the
    # manufacturer nameplate in the report, not the MW / PF estimate.
    ac_output = {
        "num_blocks": 20,
        "pcs_per_block": 4,
        "pcs_count_by_block": [4] * 20,
        "pcs_kw": 1250.0,
        "block_size_mw": 5.0,
        "selected_ratio": "1:4",
        "grid_power_factor": 0.9,
        # bound product overrides:
        "transformer_kva": 5000.0,
        "transformer_mva": 5.0,
        "transformer_vector_group": "Dyn11",
        "ac_block_product_block_code": "KEHUA-BCS5000K-C-HUD-T4",
    }
    ctx = build_report_context(
        session_state={},
        stage_outputs={
            "stage13_output": {
                "project_name": "Trunk Bind", "poi_power_req_mw": 100.0,
                "poi_energy_req_mwh": 400.0, "project_life_years": 20,
                "poi_guarantee_year": 0, "cycles_per_year": 365,
            },
            "ac_output": ac_output, "stage2": {"container_count": 80},
        },
        project_inputs={"grid_power_factor": 0.9},
    )
    # report_context takes the real transformer kva, not the pf fallback.
    assert ctx.transformer_rating_kva == 5000.0
    doc = Document(io.BytesIO(export_report_v2_1(ctx)))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            parts.extend(c.text for c in r.cells)
    text = "\n".join(parts)
    assert "Manufacturer nameplate rating" in text
    assert "÷ 0.9" not in text and "/ 0.9" not in text
