"""Two AC alternatives must export as two reports, not one overwriting the other.

Owner ruling B (2026-08-04): "同一个确定了的 DC 方案，AC 是可以稍微有多一个方案的
… 最终报告可以重新生成一个版本". Step 2 already makes the drawings follow the
selected alternative; this closes the last hole — both versions previously landed
on the SAME download name, so the second silently replaced the first.

The rule is asymmetric on purpose: a label appears ONLY when the DC run carries
more than one alternative. One alternative is the normal case and distinguishes
itself from nothing, so those file names must stay exactly as they were.
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from calb_sizing_tool.infra.db.base import Base
from calb_sizing_tool.infra.db.models.project import Project
from calb_sizing_tool.infra.db.models.sizing_run import SizingRun
from calb_sizing_tool.infra.db.session import session_scope
from calb_sizing_tool.reporting.export_docx import make_proposal_filename
from calb_sizing_tool.reporting.report_context import build_report_context
from calb_sizing_tool.reporting.report_v2 import export_report_v2_1


# --------------------------------------------------------------------------
# The file name
# --------------------------------------------------------------------------

def test_two_alternatives_download_as_two_files():
    a = make_proposal_filename("Ningxia", version="V2.1", ac_alternative="A")
    b = make_proposal_filename("Ningxia", version="V2.1", ac_alternative="B")
    assert a != b
    assert a.endswith("_V2.1_AC-A.docx")
    assert b.endswith("_V2.1_AC-B.docx")


def test_a_single_alternative_keeps_the_historical_name():
    """No label -> byte-identical to what the tool has always produced."""
    plain = make_proposal_filename("Ningxia", version="V2.1")
    assert plain == make_proposal_filename("Ningxia", version="V2.1", ac_alternative=None)
    assert plain.endswith("_V2.1.docx")
    assert "_AC-" not in plain


def test_the_alternative_label_cannot_escape_the_file_name():
    name = make_proposal_filename("P", ac_alternative="../../etc/passwd")
    assert "/" not in name and "\\" not in name and ".." not in name


def test_the_brand_prefix_and_version_still_apply_with_a_label():
    name = make_proposal_filename("P", version="V3", prefix="GX", ac_alternative="B")
    assert name.startswith("GX_P_BESS_Proposal_")
    assert name.endswith("_V3_AC-B.docx")


# --------------------------------------------------------------------------
# The label itself
# --------------------------------------------------------------------------

@pytest.fixture()
def dc_run_with(tmp_path, monkeypatch):
    """Factory: a DC run with ``n`` AC alternatives, on an isolated database."""
    url = f"sqlite:///{(tmp_path / 'alt.sqlite').as_posix()}"
    monkeypatch.setenv("CALB_DATABASE_URL", url)
    with session_scope(url) as session:
        Base.metadata.create_all(bind=session.get_bind())

    def _make(n: int) -> tuple[str, list[str]]:
        children = [f"ac-{i}" for i in range(n)]
        with session_scope(url) as session:
            session.add(Project(project_id="p1", project_code="P1", project_name="P1"))
            session.flush()
            session.add(SizingRun(
                sizing_run_id="dc-1", project_id="p1", sizing_case_id=None,
                run_type="dc_pipeline", status="succeeded",
                input_summary_json={}, output_summary_json={},
            ))
            session.flush()
            for child in children:
                session.add(SizingRun(
                    sizing_run_id=child, project_id="p1", sizing_case_id=None,
                    parent_run_id="dc-1", run_type="ac_sizing", status="succeeded",
                    input_summary_json={}, output_summary_json={},
                ))
                session.flush()
        return "dc-1", children

    return _make


def test_alternatives_are_named_oldest_first(dc_run_with):
    from calb_sizing_tool.services.ac_run_service import ac_alternative_label

    dc, children = dc_run_with(3)
    labels = [ac_alternative_label(dc, child) for child in children]
    assert labels == ["A", "B", "C"], (
        "the first alternative tried must stay A — a later one renaming it would "
        "make an already-issued report unreproducible"
    )


def test_a_lone_alternative_is_not_named(dc_run_with):
    from calb_sizing_tool.services.ac_run_service import ac_alternative_label

    dc, children = dc_run_with(1)
    assert ac_alternative_label(dc, children[0]) is None


# --------------------------------------------------------------------------
# End to end: state -> context -> document
# --------------------------------------------------------------------------

def _ctx(state: dict):
    return build_report_context(
        session_state=state,
        stage_outputs={
            "stage13_output": {
                "project_name": "Alt Project",
                "poi_power_req_mw": 40.0,
                "poi_energy_req_mwh": 160.0,
                "project_life_years": 20,
                "poi_guarantee_year": 0,
                "cycles_per_year": 365,
            },
            "ac_output": {"num_blocks": 4, "block_size_mw": 10.0, "pcs_per_block": 8},
            "stage2": {"container_count": 32},
        },
        project_inputs={"grid_power_factor": 0.9},
    )


def _text(ctx) -> str:
    doc = Document(io.BytesIO(export_report_v2_1(ctx)))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_the_selected_alternative_reaches_the_context(dc_run_with):
    dc, children = dc_run_with(2)
    ctx = _ctx({"active_run_id": dc, "active_ac_run_id": children[1]})
    assert ctx.ac_alternative_id == children[1]
    assert ctx.ac_alternative_label == "B"


def test_the_alternative_is_named_in_the_document(dc_run_with):
    dc, children = dc_run_with(2)
    text = _text(_ctx({"active_run_id": dc, "active_ac_run_id": children[1]}))
    assert "AC Alternative: B" in text, "the cover must say which alternative this is"
    assert "AC Alternative" in text and children[1] in text, (
        "provenance must record the AC run the drawings came from"
    )


def test_an_ordinary_report_gains_no_alternative_wording(dc_run_with):
    dc, children = dc_run_with(1)
    text = _text(_ctx({"active_run_id": dc, "active_ac_run_id": children[0]}))
    assert "AC Alternative" not in text


def test_the_export_page_passes_the_label_to_the_file_name():
    """Otherwise the document says B while the file downloads as the same name."""
    import inspect

    from calb_sizing_tool.ui import report_export_view

    src = inspect.getsource(report_export_view)
    assert "ac_alternative=ctx.ac_alternative_label" in src
