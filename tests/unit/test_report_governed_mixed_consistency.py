"""Report-layer consistency for a MIXED governed (Phase B) site.

Locks the fixes for the reported discrepancy where a non-multiple-of-8 project
(92 DC) exported the generic path — 23 x 4-PCS blocks + a fabricated 5.56 MVA
(= 5.0 MW / 0.9) transformer — instead of the governed decomposition. With the
governed primary output the report must:
  (1) report the TRUE site rollup (12 governed AC Blocks / 115 MW / 92 PCS), not
      the uniform head group (11) nor an average DC-per-AC reconstruction;
  (2) never fabricate a transformer nameplate from MW / PF for a governed run;
  (3) carry the governed identity so the governed report sections activate.
"""
from __future__ import annotations

from calb_sizing_tool.reporting.report_context import build_report_context
from calb_sizing_tool.reporting.report_v2 import _compute_site_layout
from calb_sizing_tool.services.governed_ac_block_service import (
    build_governed_primary_ac_output,
)


def _stage13() -> dict:
    return {
        "project_name": "Riverside BESS 115MW",
        "poi_power_req_mw": 115.0,
        "poi_energy_req_mwh": 400.0,
        "project_life_years": 20,
        "poi_guarantee_year": 0,
        "cycles_per_year": 365,
    }


def _mixed_ctx():
    ac_output = build_governed_primary_ac_output(92)  # 11x8 + 1x4, settings-only
    return build_report_context(
        session_state={},
        stage_outputs={
            "stage13_output": _stage13(),
            "ac_output": ac_output,
            "stage2": {"container_count": 92},
        },
        project_inputs={},
    )


def test_report_uses_true_governed_site_rollup_not_head_only():
    ctx = _mixed_ctx()
    assert ctx.configuration_code == "ACBLK-10MW-8PCS-8DC-40FT-BILATERAL"
    # Site rollup, not the 11-block uniform head.
    assert ctx.ac_blocks_total == 12
    assert ctx.pcs_modules_total == 92
    assert ctx.dc_blocks_total == 92


def test_report_never_fabricates_transformer_mva_for_governed_mixed():
    ctx = _mixed_ctx()
    # No product / no owner MVA -> TBD, never 5.0/0.9 = 5.56 MVA or 10/0.9.
    assert ctx.transformer_rating_kva is None


def test_governed_groups_carried_for_the_report():
    ctx = _mixed_ctx()
    groups = ctx.ac_output.get("governed_groups")
    assert isinstance(groups, list) and len(groups) == 2
    assert {g["configuration_code"] for g in groups} == {
        "ACBLK-10MW-8PCS-8DC-40FT-BILATERAL",
        "ACBLK-5MW-4PCS-4DC-40FT-LINEAR",
    }


def test_linear_site_array_suppressed_for_governed_mixed():
    ctx = _mixed_ctx()
    # A governed run shows the governed composition (§9), not the generic linear
    # L2 site array.
    assert _compute_site_layout(ctx) is None


def test_full_report_doc_builds_for_governed_mixed():
    # The whole DOCX must assemble without error and reflect the governed site.
    import io

    from docx import Document

    from calb_sizing_tool.reporting.report_v2 import export_report_v2_1

    ctx = _mixed_ctx()
    doc_bytes = export_report_v2_1(ctx)
    doc = Document(io.BytesIO(doc_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(parts)
    # The true mixed site is stated (both product models, 115 MW total); the
    # generic 4-per-block average and the MW/PF fabrication are absent.
    assert "ACBLK-10MW-8PCS-8DC-40FT-BILATERAL" in text
    assert "ACBLK-5MW-4PCS-4DC-40FT-LINEAR" in text
    assert "115" in text  # true site MW appears somewhere
    assert "÷ 0.9" not in text and "/ 0.9" not in text


def test_customer_report_has_no_internal_jargon_or_code_identifiers():
    """The exported report is customer-facing: it must not leak internal code
    module names, developer jargon or pipeline stage-names."""
    import io

    from docx import Document

    from calb_sizing_tool.reporting.report_v2 import export_report_v2_1

    ctx = _mixed_ctx()
    doc = Document(io.BytesIO(export_report_v2_1(ctx)))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(parts).lower()
    for token in (
        "ac_power_reconciliation", "phase b", "phase a", "governed", "layout_variant",
        "mw ÷ pf", "owner-confirmed", "owner_selected", "provisional_auto", "rule a",
        "read-only", "settings-only", "frozen stage", "l3 master", "not-for-construction",
    ):
        assert token not in text, f"internal token leaked into customer report: {token!r}"
