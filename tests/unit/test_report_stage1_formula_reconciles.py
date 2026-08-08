"""Section 3 must reconcile: its numbers, in its formula, give its answer.

A customer who checks the arithmetic is the point of stating a formula at all.
Section 3 used to state

    DC Energy Required = POI ÷ ((1 − S&C) × DoD × DC RTE (discharge) × One-way Eff)

and then print "DC RTE: 94.00%" — the ROUND-TRIP figure. The denominator takes
the root of it (only the discharge half of a round trip delivers energy to the
POI), so a reader following the document got 472.49 MWh where the report said
458.09 MWh: 14.4 MWh, 3.1%, in a document whose purpose is to be checked.

The sizing formula was correct and is untouched (frozen canon, and owner ruling
2026-08-08 "DC SIZING 等公式坚决不能动"). Only the report's statement of it
changed. This test reads the numbers back out of the generated document and
recomputes, so the two can never drift apart again.
"""
from __future__ import annotations

import io
import math
import re
from pathlib import Path

import pytest

pytest.importorskip("docx")
from docx import Document

from calb_sizing_tool.config import DC_DATA_PATH
from calb_sizing_tool.reporting.report_context import build_report_context
from calb_sizing_tool.reporting.report_v2 import export_report_v2_1
from calb_sizing_tool.ui import dc_view


@pytest.fixture(scope="module")
def report_paragraphs_and_tables(tmp_path_factory, ) -> tuple[list[str], list[list[list[str]]]]:
    defaults, *_ = dc_view.load_data(Path(DC_DATA_PATH))
    stage1 = dc_view.run_stage1(
        {
            "project_name": "Formula Reconciles",
            "poi_power_req_mw": 100.0,
            "poi_energy_req_mwh": 400.0,
            "eff_dc_cables": 99.5, "eff_pcs": 98.5, "eff_mvt": 99.5,
            "eff_ac_cables_sw_rmu": 99.2, "eff_hvt_others": 100.0,
            "sc_time_months": 3, "sc_loss_frac": dc_view.calc_sc_loss_pct(3) / 100.0,
            "dod_pct": 95.0, "dc_round_trip_efficiency_pct": 94.0,
            "rte_curve_adjust_pp": 0.0, "rte_monotonic_enforce": True,
            "project_life_years": 20, "cycles_per_year": 365, "poi_guarantee_year": 10,
        },
        defaults,
    )
    ctx = build_report_context(
        session_state={},
        stage_outputs={
            "stage13_output": stage1,
            "stage2": {"container_count": 80},
            "ac_output": {"num_blocks": 20, "block_size_mw": 5.0, "pcs_per_block": 2},
        },
        project_inputs={},
    )
    doc = Document(io.BytesIO(export_report_v2_1(ctx)))
    paragraphs = [p.text for p in doc.paragraphs]
    tables = [[[c.text for c in row.cells] for row in t.rows] for t in doc.tables]
    return paragraphs, tables


def _percent_after(paragraphs: list[str], label: str) -> float:
    """The percentage the document prints against `label`, as a fraction."""
    for text in paragraphs:
        if label in text:
            tail = text.split(label, 1)[1]
            match = re.search(r"([\d.]+)\s*%", tail)
            if match:
                return float(match.group(1)) / 100.0
    raise AssertionError(f"the report no longer prints {label!r}")


def _table_value(tables, metric: str) -> float:
    for table in tables:
        for row in table:
            if row and row[0].strip() == metric:
                return float(re.sub(r"[^\d.\-]", "", row[1]))
    raise AssertionError(f"the report no longer prints the {metric!r} row")


def test_the_printed_numbers_reproduce_the_printed_answer(report_paragraphs_and_tables):
    """The whole point. Every input comes from the DOCX, not from the code."""
    paragraphs, tables = report_paragraphs_and_tables

    sc = _percent_after(paragraphs, "S&C loss:")
    dod = _percent_after(paragraphs, "DoD:")
    discharge = _percent_after(paragraphs, "DC RTE (one-way discharge):")
    chain = _percent_after(paragraphs, "One-way Efficiency (DC→POI):")

    poi = _table_value(tables, "POI Energy Requirement (MWh)")
    printed = _table_value(tables, "DC Energy Capacity Required (MWh)")

    recomputed = poi / ((1 - sc) * dod * discharge * chain)

    assert recomputed == pytest.approx(printed, abs=0.05), (
        f"a reader applying section 3's formula to section 3's numbers gets "
        f"{recomputed:.2f} MWh, but the report states {printed:.2f} MWh"
    )


def test_the_document_says_which_rte_the_denominator_takes(report_paragraphs_and_tables):
    """Both figures appear, and the root is explicit rather than implied."""
    paragraphs, _tables = report_paragraphs_and_tables
    joined = "\n".join(paragraphs)

    assert "DC RTE (one-way discharge)" in joined
    assert "DC RTE (round-trip)" in joined
    assert "√(DC RTE round-trip)" in joined

    round_trip = _percent_after(paragraphs, "DC RTE (round-trip):")
    discharge = _percent_after(paragraphs, "DC RTE (one-way discharge):")
    assert discharge == pytest.approx(math.sqrt(round_trip), abs=0.0001)
    assert discharge > round_trip, "the discharge half is the larger fraction"


def test_the_formula_no_longer_reads_as_the_round_trip_value(report_paragraphs_and_tables):
    """The exact wording that misled: "DC RTE (discharge)" beside a round-trip number."""
    paragraphs, _tables = report_paragraphs_and_tables
    for text in paragraphs:
        if text.startswith("DC Energy Required (MWh) ="):
            assert "× DC RTE (discharge) ×" not in text, (
                "the formula names a discharge figure the document does not print"
            )
            return
    raise AssertionError("section 3 no longer states the formula at all")
