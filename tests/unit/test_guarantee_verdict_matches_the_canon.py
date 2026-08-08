"""One guarantee verdict, and it is the canon's.

SIZING_LOGIC_CANON_V1 §9 ("Guarantee-year 扩容规则不可变") fixes the test:

    若 guarantee year 的 poi_usable_energy_mwh + 1e-6 < poi_energy_req_mwh，
    则继续加设备

`dc_pipeline_service` applies it verbatim to decide `converged`, and the DC page,
the run registry and the workbench all just display that flag. In the product
the report is comparing the same number against the same target —
`report_export_view` sets `poi_energy_guarantee_mwh` to
`stage13_output["poi_energy_req_mwh"]` — so every surface should agree.

Section 1 of the report used to allow a 0.1 MWh shortfall. Nothing justified it:
not the canon, not the report specs, not the git history (both criteria entered
with the repository's squashed initial import). It made section 1 the only one
of six verdicts that could pass a design the engine had failed, and it
contradicted section 5's own table inside the same document.
"""
from __future__ import annotations

import inspect
import re

import pandas as pd
import pytest

from calb_sizing_tool.reporting import report_v2
from calb_sizing_tool.reporting.report_v2 import GUARANTEE_EPSILON_MWH
from calb_sizing_tool.services import dc_pipeline_service


def test_the_report_uses_the_epsilon_the_engine_uses():
    """The engine holds the value inline; this is what keeps the copy honest."""
    source = inspect.getsource(dc_pipeline_service.size_with_guarantee)
    epsilons = set(re.findall(r"poi_g \+ ([0-9.e\-]+) >= poi_energy_req", source))

    assert epsilons, "the engine's convergence test no longer looks like the canon's"
    assert len(epsilons) == 1, f"the engine uses more than one epsilon: {epsilons}"
    assert float(epsilons.pop()) == GUARANTEE_EPSILON_MWH


def test_the_epsilon_is_a_float_guard_not_an_engineering_allowance():
    """0.1 MWh would be an allowance. 1e-6 cannot change a real verdict."""
    assert GUARANTEE_EPSILON_MWH <= 1e-6


def _ctx(usable: float, target: float):
    from calb_sizing_tool.reporting.report_context import build_report_context

    ctx = build_report_context(
        session_state={},
        stage_outputs={
            "stage13_output": {
                "project_name": "Verdict", "poi_power_req_mw": 100.0,
                "poi_energy_req_mwh": target, "project_life_years": 20,
                "poi_guarantee_year": 10, "cycles_per_year": 365,
            },
            "stage2": {"container_count": 80},
            "ac_output": {"num_blocks": 20, "block_size_mw": 5.0, "pcs_per_block": 2},
            "stage3_df": pd.DataFrame({
                "Year_Index": [0, 10],
                "POI_Usable_Energy_MWh": [target + 20.0, usable],
                "DC_Usable_MWh": [target + 30.0] * 2,
                "SOH_Display_Pct": [100.0] * 2, "SOH_Absolute_Pct": [100.0] * 2,
                "DC_RTE_Pct": [94.0] * 2, "System_RTE_Pct": [92.0] * 2,
            }),
        },
        project_inputs={"poi_energy_guarantee_mwh": target},
    )
    return ctx


def _verdicts(ctx) -> tuple[str, str]:
    """(section 1's verdict, section 5's guarantee-year row verdict) from a real DOCX."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(report_v2.export_report_v2_1(ctx)))
    section1 = None
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            if cells and cells[0].strip() == "Guarantee Compliance":
                section1 = cells[1].strip()
    section5 = None
    for table in doc.tables:
        header = [c.text for c in table.rows[0].cells]
        matches = [i for i, h in enumerate(header) if h.startswith("Meets")]
        if not matches:
            continue
        index = matches[0]
        for row in table.rows[1:]:
            value = row.cells[index].text.strip()
            if value in ("Yes", "No"):
                section5 = value
    return section1, section5


@pytest.mark.parametrize(
    "usable, target, expected",
    [
        (400.00, 400.00, "Yes"),   # exactly on target
        (400.50, 400.00, "Yes"),   # comfortably over
        (399.95, 400.00, "No"),    # the band the 0.1 slack used to pass
        (399.99, 400.00, "No"),
        (350.00, 400.00, "No"),
    ],
)
def test_both_sections_give_the_same_verdict(usable, target, expected):
    """The contradiction that made this a defect regardless of which test is right."""
    section1, section5 = _verdicts(_ctx(usable, target))

    assert section1 == expected, f"section 1 said {section1} for {usable} vs {target}"
    assert section5 == expected, f"section 5 said {section5} for {usable} vs {target}"
    assert section1 == section5


def test_the_shortfall_the_slack_used_to_hide_now_reads_as_a_miss():
    """The specific figures from the 2026-08-08 reproduction."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(report_v2.export_report_v2_1(_ctx(399.95, 400.00))))
    summary = [p.text for p in doc.paragraphs if "delivers" in p.text and "guarantee" in p.text]

    assert summary, "section 1 no longer states the headline sentence"
    assert "399.95" in summary[0] and "400.00" in summary[0]
    assert "guarantee NOT met" in summary[0], (
        f"the numbers printed in this very sentence show a miss: {summary[0]}"
    )
